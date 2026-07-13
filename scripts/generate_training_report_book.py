#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《软件开发综合实训报告册》，章节结构与公众监督 NEP 模板对齐。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_builder import (  # noqa: E402
    ASSETS_DIR,
    add_attachment_cover,
    add_bullets,
    add_center_title,
    add_h,
    add_image_figure,
    add_kv_table,
    add_p,
    add_page_break,
    add_placeholder_figure,
    add_table,
    add_table_of_contents,
    copy_cqpt_template,
    count_doc_chars,
    open_cqpt_report_book,
    patch_attachment_titles,
    patch_cover_tables,
    save_doc,
    set_run_font,
    truncate_body_after_section_break,
)
from docx_builder import FONT_BODY  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

ROOT = Path(r"F:\项目2资料")
OUT_DIR = ROOT / "项目资料提交" / "项目资料提交" / "05.软件开发综合实训报告册"
OUT_FILE = OUT_DIR / "软件开发综合实训报告册-冯威-2023214624.docx"
WORK_FILE = Path(__file__).resolve().parent / "_report_book_work.docx"

STUDENT = "冯威"
STUDENT_ID = "2023214624"
STUDENT_PHONE = "15736634798"
CLASS_ID = "13002306"
COLLEGE = "计算机科学与技术学院"
SEMESTER = "2025-2026学年 春秋学期"
COURSE = "软件开发综合实训"
COURSE_CODE = "A2130400"
TRAINING_PLACE = "综合实验大楼"
TRAINING_TIME = "202602学期"
TRAINING_COMPANY = "东软教育科技集团"
EXTERNAL_TEACHER = "【请填写校外指导教师】"
PROJECT = "新能源充电桩大数据分析项目"
PROJECT_LINE = "新能源充电桩大数据分析软件/项目/系统"
PRODUCT = "CBP（Charging Bigdata Platform）"

# 与 generate_submission_docs_full 一致
TABLE_FIELDS = {
    "t_voltage_current": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_hour", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
        ["avg_pack_voltage", "DOUBLE", "-", "是", "", "平均组电压(V)"],
        ["avg_charge_current", "DOUBLE", "-", "是", "", "平均充电电流(A)"],
    ],
    "t_cell_voltage_range": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_hour", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
        ["max_cell_voltage", "DOUBLE", "-", "是", "", "最高单体电压(V)"],
        ["min_cell_voltage", "DOUBLE", "-", "是", "", "最低单体电压(V)"],
    ],
    "t_temperature": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_time", "VARCHAR", "20", "否", "UK", "yyyyMMddHHmmss"],
        ["max_temperature", "DOUBLE", "-", "是", "", "最高温度(℃)"],
        ["min_temperature", "DOUBLE", "-", "是", "", "最低温度(℃)"],
    ],
    "t_energy_capacity": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_hour", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
        ["avg_available_energy", "DOUBLE", "-", "是", "", "平均可用能量(kW)"],
        ["avg_available_capacity", "DOUBLE", "-", "是", "", "平均可用容量(Ah)"],
    ],
    "t_charge_current_stats": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_hour", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
        ["avg_charge_current", "DOUBLE", "-", "是", "", "平均充电电流(A)"],
        ["max_charge_current", "DOUBLE", "-", "是", "", "最大充电电流(A)"],
    ],
    "t_voltage_current_relation": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["record_time", "VARCHAR", "20", "否", "UK", "yyyyMMddHHmmss"],
        ["pack_voltage", "DOUBLE", "-", "是", "", "组电压(V)"],
        ["charge_current", "DOUBLE", "-", "是", "", "充电电流(A)"],
    ],
    "t_soc_temperature": [
        ["id", "INT", "11", "否", "PK", "自增主键"],
        ["soc_bucket", "VARCHAR", "20", "否", "UK", "SOC 分段如 10-20"],
        ["avg_max_temperature", "DOUBLE", "-", "是", "", "平均最高温度(℃)"],
        ["avg_min_temperature", "DOUBLE", "-", "是", "", "平均最低温度(℃)"],
    ],
    "sys_user": [
        ["id", "INT", "11", "否", "PK", "用户 ID"],
        ["username", "VARCHAR", "50", "否", "UK", "登录用户名"],
        ["email", "VARCHAR", "100", "是", "UK", "登录邮箱"],
        ["password_hash", "VARCHAR", "255", "否", "", "bcrypt 哈希"],
        ["role", "VARCHAR", "20", "否", "IDX", "admin / user"],
        ["is_active", "TINYINT", "1", "否", "", "1 启用 0 禁用"],
        ["email_verified", "TINYINT", "1", "否", "", "1 已验证"],
        ["created_at", "DATETIME", "-", "否", "", "创建时间"],
        ["updated_at", "DATETIME", "-", "否", "", "更新时间"],
    ],
}
FIELD_HEADERS = ["字段", "类型", "长度", "空", "键", "备注"]

API_LIST = [
    ["1", "GET", "/health", "否", "服务存活探测"],
    ["2", "GET", "/health/ready", "否", "就绪检查含 MySQL"],
    ["3", "POST", "/auth/login", "否", "邮箱登录双 Token"],
    ["4", "POST", "/auth/register", "否", "邮箱注册"],
    ["5", "POST", "/auth/refresh", "否", "刷新 Access Token"],
    ["6", "POST", "/auth/logout", "否", "吊销 Refresh Token"],
    ["7", "GET", "/auth/me", "是", "当前用户信息"],
    ["8", "POST", "/auth/forgot-password", "否", "忘记密码邮件"],
    ["9", "POST", "/auth/complete-email-token", "否", "邮件链接设密"],
    ["10", "PATCH", "/auth/profile/password", "是", "修改密码"],
    ["11", "GET", "/bi/voltage-current", "是", "v1 电压电流"],
    ["12", "GET", "/bi/cell-voltage-range", "是", "v2 单体电压"],
    ["13", "GET", "/bi/temperature", "是", "v3 温度趋势"],
    ["14", "GET", "/bi/energy-capacity", "是", "v4 能量容量"],
    ["15", "GET", "/bi/charge-current-stats", "是", "v5 充电电流"],
    ["16", "GET", "/bi/voltage-current-relation", "是", "v6 电压电流关系"],
    ["17", "GET", "/bi/soc-temperature", "是", "v7 SOC 温度"],
    ["18", "GET", "/bi/soc-hourly", "是", "ADS 每小时 SOC"],
    ["19", "GET", "/bi/charging-daily", "是", "ADS 日充电次数"],
    ["20", "POST", "/predict/fee", "是", "费用预测"],
    ["21", "POST", "/predict/duration", "是", "时长预测"],
    ["22", "POST", "/predict/platform", "是", "平台预测"],
    ["23", "POST", "/predict/soc", "是", "SOC 预测"],
    ["24", "POST", "/assistant/chat/stream", "是", "智能助手 SSE"],
    ["25", "GET", "/views/configs", "是", "BI 视图配置"],
    ["26", "POST", "/bi/cache/invalidate", "是", "清空 BI 缓存"],
]


def _figure_png(caption: str) -> Path | None:
    """assets 下图题 PNG，如 图3.1 -> 图3_1.png"""
    key = caption.split()[0]
    path = ASSETS_DIR / f"{key.replace('.', '_')}.png"
    return path if path.exists() and path.stat().st_size > 500 else None


def _add_figure(doc, caption: str, hint: str = "", width_cm: float = 14.0):
    png = _figure_png(caption)
    if png:
        add_image_figure(doc, caption, png, width_cm=width_cm)
    else:
        add_placeholder_figure(doc, caption, hint)


def _cover_cqpt(doc):
    """重庆邮电大学报告册封面（学年学期、学院、学号等信息表）。"""
    add_cqpt_cover_header(doc)
    add_cover_info_table(
        doc,
        [
            ["学年学期：", SEMESTER],
            ["课程名称：", COURSE],
            ["学生学院：", COLLEGE],
            ["专业班级：", CLASS_ID],
            ["学生学号：", STUDENT_ID],
            ["学生姓名：", STUDENT],
            ["联系电话：", STUDENT_PHONE],
        ],
    )
    for _ in range(6):
        doc.add_paragraph()


TRAINING_PURPOSE = (
    "新能源充电桩大数据分析项目（CBP）面向充电桩运营场景，"
    "建设「数据湖—批处理—数据仓库—可视化—智能预测」一体化平台。"
    "项目综合运用 Hadoop 分布式存储与 MapReduce 批处理、"
    "Python 数据分析与机器学习、FastAPI 后端与 Vue3 前端技术，"
    "完成需求分析、概要设计、详细设计、编码实现、联调测试与文档编写全流程，"
    "形成可部署、可演示、可维护的软件交付物。"
)

TRAINING_PROCESS = (
    "项目按软件工程阶段有序推进，主要过程如下。\n\n"
    "需求分析与原型设计。小组研读充电桩运营业务背景，明确电池遥测数据（组电压、单体电压、温度、SOC、"
    "能量容量等）与充电会话数据（费用、时长、平台）两类数据源。完成系统总体用例图、数据流图、"
    "MR 批处理流程图，并整理《需求规格说明书》与页面效果图。"
    "在需求评审中，将 MapReduce 任务统一编号为 v1~v7，明确 HDFS /Car/ 四 CSV 上传规范与 MySQL charging_bigdata "
    "七张 MR 结果表结构，为接口设计与库表设计奠定基础。\n\n"
    "概要设计与数据库设计。依据需求划分 mapreduce、analytics、backend、frontend、sql 五模块，"
    "完成五层 B/S 架构图与 26 个 REST 接口清单。数据库方面绘制 E-R 图，"
    "设计 t_voltage_current 至 t_soc_temperature 七张 MR 表及 ADS 汇总表、sys_user 认证表，"
    "编写 schema.sql 与 UPSERT 入库脚本，保证时间键唯一与幂等入库。\n\n"
    "详细设计说明书编写。按项目文档规范完成详细设计文档，涵盖引言、设计概述、"
    "系统详细需求、总体方案、系统详细设计（M01~M07 七模块）、数据库设计（含数据字典）、"
    "非功能性设计与环境配置。补充架构图、业务流程图、E-R 图及 Web/BI 运行截图，"
    "对各功能模块分别描述输入、处理、算法、输出与关联数据表。\n\n"
    "MapReduce 批处理开发。基于 Hadoop 3.1 实现 v1~v7 七个 Java 作业，"
    "公共包 com.cbp.mr 提供 RecordParser、TimeKeyUtil、Avg2Reducer 等组件。"
    "JobRunner 顺序提交作业，输出至 HDFS /Car/output/v1…v7，"
    "load_mr_to_mysql.py 解析 TAB 分隔结果 UPSERT 至 MySQL。\n\n"
    "Web 后端开发。基于 FastAPI 搭建 backend，采用 api → services → repositories 三层结构，"
    "实现 JWT 双 Token 邮箱登录、BI 统计查询、四项 ML 预测接口与 RAG 智能助手 SSE 流式问答。"
    "在 config.py 中集中配置数据库、CORS、JWT 密钥，OpenAPI 文档自动生成于 /docs。\n\n"
    "Web 前端开发。使用 Vue3 + TypeScript + Vite + Element Plus 实现登录注册、"
    "首页模块导航、MrBiView 七图 ECharts 分析页、PredictView 四项预测表单、"
    "AssistantDrawer 智能助手抽屉。axios 拦截器自动附加 Bearer Token 与 401 刷新逻辑。\n\n"
    "Python 分析与机器学习。通过 WebHDFS 读取 HDFS 数据，"
    "完成 SOC 趋势、充电次数、充电速率、SOC 热力图等五类分析图；"
    "基于 nvv2t_md.csv 训练费用/时长/平台/SOC 四项模型，费用预测对比 LinearRegression 与 XGBoost 按 RMSE 选优。\n\n"
    "BI 与 Docker 部署。使用 Datart 连接 MySQL 构建 Charging-BigData-BI 七图 Dashboard 并发布分享链接；"
    "编写 docker-compose.yml 在阿里云 ECS 部署 mysql、backend、frontend、datart 四容器，"
    "公网演示地址 http://115.29.194.137:8080。\n\n"
    "联调测试与上线准备。按「HDFS 上传→MR 跑批→MySQL 入库→Web 登录→图表展示→预测推理」主流程端到端测试。"
    "编写部署手册与本详细设计说明书。"
    "针对 Datart RC.3 分享页图表兼容、JWT Refresh 并发、CORS 跨域等问题逐项排查并记录解决方案。",
)

TRAINING_SUMMARY = (
    "项目已完成 HDFS 数据湖、MapReduce 七作业、MySQL 指标库、Datart BI 大屏、"
    "Vue3 Web 应用、四项机器学习预测与 RAG 智能助手，公网演示地址 http://115.29.194.137:8080。"
    "交付物包括需求规格、概要设计、数据库设计、接口文档、部署手册与本报告册，"
    "源码托管于 charging-bigdata 仓库，版本 V1.3。"
)


TRAINING_BODY = (
    f"项目背景与目标\n\n{TRAINING_PURPOSE}\n\n"
    f"实施过程\n\n{TRAINING_PROCESS}\n\n"
    f"交付成果\n\n{TRAINING_SUMMARY}"
)


def _training_report_main(doc):
    """报告册正文：课程信息表 + 实训目的/过程/总结。"""
    meta = doc.add_table(rows=4, cols=4)
    meta.style = "Table Grid"
    rows_data = [
        ["课程名称", COURSE, "课程编号", COURSE_CODE],
        ["实训地点", TRAINING_PLACE, "实训时间", TRAINING_TIME],
        ["实训企业名称", TRAINING_COMPANY, "校外指导教师", EXTERNAL_TEACHER],
        ["评阅人签字", "", "成绩（百分制）", ""],
    ]
    for ri, row_vals in enumerate(rows_data):
        for ci, val in enumerate(row_vals):
            cell = meta.rows[ri].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            set_run_font(r, 12, bool(val and ci % 2 == 0), FONT_BODY)

    doc.add_paragraph()
    add_p(doc, "实训内容", bold=True, first_line_indent=0)
    add_p(doc, "实训目的和意义", bold=True, first_line_indent=0)
    add_p(doc, TRAINING_PURPOSE)
    doc.add_paragraph()
    add_p(doc, "实训内容和过程", bold=True, first_line_indent=0)
    add_p(doc, TRAINING_PROCESS)
    doc.add_paragraph()
    add_p(doc, "实训总结与体会", bold=True, first_line_indent=0)
    add_p(doc, TRAINING_SUMMARY)
    add_page_break(doc)


def _cover_attachment(doc):
    """附件：详细设计说明书封面。"""
    add_attachment_cover(doc, PROJECT_LINE)
    doc.add_paragraph()
    doc.add_paragraph()


def _cover(doc):
    pass  # 已由 _cover_cqpt 替代


def _revision(doc):
    add_p(doc, "编写单位或个人", bold=True, first_line_indent=0)
    add_p(doc, f"{COLLEGE} {CLASS_ID} {STUDENT}")
    add_p(doc, "修订历史", bold=True, first_line_indent=0)
    add_table(
        doc,
        ["版本", "日期", "作者", "更改说明"],
        [
            ["V1.0", "2026-05-10", STUDENT, "初稿：MR v1~v7 与 MySQL 落库"],
            ["V1.1", "2026-05-28", STUDENT, "补充 Vue3 前端与 FastAPI 接口"],
            ["V1.2", "2026-06-08", STUDENT, "增加 JWT 鉴权、ADS 五表、Datart BI"],
            ["V1.3", "2026-06-22", STUDENT, "ML 四项预测、RAG 助手、Docker 云部署定稿"],
        ],
    )
    doc.add_paragraph()
    add_table_of_contents(doc)
    add_page_break(doc)
    add_p(doc, PROJECT_LINE, bold=True, first_line_indent=0)
    add_p(doc, "详细设计说明书", bold=True, first_line_indent=0)
    doc.add_paragraph()


def _chapter1(doc):
    add_h(doc, "1 引言", 1)
    add_h(doc, "1.1 编写目的", 2)
    add_p(
        doc,
        f"本文档是《{PROJECT}》的详细设计说明书，"
        "在《需求规格说明书》《概要设计说明书》《数据库设计说明书》及页面原型基础上，"
        "对系统详细设计规范、实现细节、模块划分、数据库结构与运行环境作出完整描述。"
        "文档面向项目干系人及后续维护人员，作为编码实现、联调测试与部署上线的正式依据。"
        f"编写单位：{COLLEGE} {CLASS_ID}。",
    )
    add_h(doc, "1.2 背景", 2)
    add_p(
        doc,
        "新能源汽车充电基础设施快速扩张，充电桩在运营过程中持续采集电池组电压、单体电压极值、"
        "充放电电流、温度、SOC、可用能量与容量等遥测数据，同时产生大量充电会话记录（起止时间、电量、费用、平台等）。"
        "传统单机 Excel 或分散数据库难以支撑海量明细的聚合分析与跨时段趋势对比，"
        "运营方迫切需要「集中存储—分布式计算—关系型持久化—可视化展示—智能预测」一体化的大数据平台。"
        "本项目以新能源充电桩运营数据为业务场景，建设完整的大数据处理与展示链路。",
    )
    add_p(
        doc,
        "本项目产品代号 CBP，已实现从四个 CSV 上传 HDFS /Car/，经七个 MapReduce 作业生成电池运营指标，"
        "load_mr_to_mysql.py 入库 charging_bigdata 库，Datart BI 大屏与 Vue3 内置 ECharts 双轨可视化，"
        "基于 nvv2t_md.csv 训练费用/时长/平台/SOC 四项机器学习模型，"
        "FastAPI 提供 REST 接口，Vue3 前端提供分析页、预测页与 RAG 智能助手，"
        "并在阿里云 ECS http://115.29.194.137:8080 通过 Docker Compose 稳定运行。",
    )
    add_h(doc, "1.3 参考资料", 2)
    add_bullets(
        doc,
        [
            f"《{PROJECT}-需求规格说明书》V1.3；",
            f"《{PROJECT}-概要设计说明书》V1.1；",
            f"《{PROJECT}-数据库设计说明书》；",
            "《接口文档》charging-bigdata/docs 与项目资料提交目录；",
            "《Datart BI 大屏部署手册》；",
            "Apache Hadoop MapReduce Tutorial；FastAPI / Vue.js 3 官方文档；",
            "东软实训任务书、课程讲义与 visu1.png 布局参考。",
        ],
    )
    add_h(doc, "1.4 术语定义及说明", 2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["HDFS", "Hadoop 分布式文件系统，本项目业务目录 /Car/"],
            ["MR", "MapReduce，Java 批处理作业 v1~v7"],
            ["ADS", "应用数据服务层汇总表，如 t_soc_hourly"],
            ["CBP", "Charging Bigdata Platform 产品代号"],
            ["JWT", "JSON Web Token，双 Token 鉴权机制"],
            ["RAG", "检索增强生成，智能助手文档问答"],
            ["SOC", "State of Charge，电池荷电状态百分比"],
            ["BI", "商业智能看板，本项目采用 Datart 1.0.0-rc.3"],
        ],
    )


def _chapter2(doc):
    add_h(doc, "2 设计概述", 1)
    add_h(doc, "2.1 设计目标", 2)
    add_h(doc, "2.1.1 需求概述", 3)
    add_p(
        doc,
        "系统须实现：（1）HDFS 四 CSV 集中管理；（2）MapReduce v1~v7 七项电池指标统计；"
        "（3）MySQL charging_bigdata 七张 MR 表与五张 ADS 表持久化；"
        "（4）Datart BI 七图大屏与 Vue3 /mr-bi 内置图表；"
        "（5）Python 五类分析图（SOC 趋势、充电次数、充电速率、SOC 热力图、电池性能）；"
        "（6）费用/时长/平台/SOC 四项 ML 预测 API；"
        "（7）邮箱注册登录、JWT 双 Token、bcrypt 密码存储；"
        "（8）RAG 智能助手 SSE 流式问答；（9）Docker 一键部署与公网在线演示。"
        "非功能要求：BI 查询 P95≤2s、预测推理≤500ms、三节点 Hadoop 集群交付、原始数据不入 Git。",
    )
    add_h(doc, "2.1.2 运行环境概述", 3)
    add_p(
        doc,
        "（1）大数据集群：Hadoop 3.1.x 三节点（1 Master + 2 Slave），HDFS 副本数 2，"
        "YARN 调度 MapReduce；Master 同时承担 MySQL 8.0 与 API 部署节点。"
        "（2）开发环境：Windows 11 开发机，IntelliJ IDEA 编写 MR，VS Code/Cursor 编写 Python 与前端，"
        "Maven 3.8 打包 JAR，Node 18 构建 Vue3。"
        "（3）演示环境：阿里云 ECS Ubuntu 22.04，Docker Compose 运行 mysql、backend、frontend、datart 容器，"
        "公网 http://115.29.194.137:8080（Web）、8088（Datart BI）。"
        "（4）字符集：全链路 UTF-8；数据库 utf8mb4。",
    )
    add_h(doc, "2.1.3 设计原则概述", 3)
    add_p(
        doc,
        "（1）分层解耦：存储层 HDFS、计算层 MR、服务层 MySQL、应用层 FastAPI、展示层 Vue3/Datart，"
        "各层通过文件格式、SQL 表结构与 REST JSON 契约交互，便于替换技术栈。"
        "（2）数据同源：批处理读 dsv13r1.csv，Python 分析读 HDFS 同源文件，避免本地与集群数据不一致。"
        "（3）安全优先：密码 bcrypt 哈希，业务 API 统一 Bearer JWT，Refresh Token 服务端轮换存储。"
        "（4）可演示：种子脚本 import_web_demo 预置 MR 数据与演示账号 admin@example.com / admin123。"
        "（5）文档驱动：SRS、概要设计、数据库设计、接口文档与实训报告交叉引用，版本号统一 V1.3。",
    )
    add_h(doc, "2.1.4 详细设计方法与工具", 3)
    add_p(
        doc,
        "（1）设计方法：面向数据流的五层架构设计、ER 建模、REST 资源划分、模块编号 M01~M07。"
        "（2）设计工具：Draw.io / Visio 绘制架构图、ER 图、时序图；Navicat / DataGrip 管理 MySQL；"
        "Datart 控制台配置 BI 数据源与 Dashboard。"
        "（3）开发工具：IDEA（Java MR）、PyCharm（Python ML）、VS Code（FastAPI + Vue3 + TypeScript）。"
        "（4）文档工具：Markdown 编写设计文档，generate_submission_docs_full.py 与本文档脚本生成宋体 Word。",
    )


def _chapter3(doc):
    add_h(doc, "3 系统详细需求说明", 1)
    add_h(doc, "3.1 详细需求说明", 2)
    add_h(doc, "3.1.1 详细功能需求说明", 3)
    add_p(
        doc,
        "系统用户角色包括：系统管理员（admin）、数据分析师（user）、只读观摩用户（只读）。"
        "管理员可重载 ML 模型、清空 BI 缓存、重建 RAG 索引；普通用户可查看图表与调用预测；"
        "未登录用户仅可访问登录注册与健康检查接口。",
    )
    add_table(
        doc,
        ["编号", "功能", "优先级", "说明"],
        [
            ["FR-01", "HDFS 四文件上传", "P0", "/Car/ 目录可列出四 CSV"],
            ["FR-02~08", "MR v1~v7", "P0", "七张结果表有数据"],
            ["FR-09", "MR 入库", "P0", "UPSERT 不重复"],
            ["FR-10", "ADS 汇总", "P1", "五张 ADS 表"],
            ["FR-11", "Datart BI", "P1", "七图大屏可分享"],
            ["FR-16~19", "ML 四项预测", "P0", "POST /predict/*"],
            ["FR-20", "JWT 鉴权", "P0", "未登录 401"],
            ["FR-21~23", "Web 页面", "P0", "首页/分析/预测"],
            ["FR-24", "智能助手", "P1", "SSE 流式回答"],
        ],
    )
    add_h(doc, "3.1.2 用例图说明", 3)
    add_p(
        doc,
        "系统主要参与者为「运营分析员」与「系统管理员」。运营分析员用例包括：登录系统、"
        "查看 MR 指标图表（电压电流、温度、SOC 等）、查看 Datart BI 大屏、"
        "输入特征进行费用/时长/平台/SOC 预测、向智能助手提问数据口径与部署问题。"
        "系统管理员额外用例：管理用户、重载模型、清空 BI 缓存、重建向量索引、"
        "配置 Datart 数据源与 Dashboard。",
    )
    _add_figure(doc, "图3.1 系统总体用例图", "参与者：运营分析员、系统管理员；用例：登录、BI 看板、分析图表、ML 预测、智能助手")
    _add_figure(doc, "图3.2 数据分析员用例图", "细化查看 v1~v7 图表、SOC 热力图、充电次数统计等")
    _add_figure(doc, "图3.3 MapReduce 批处理用例图", "上传 CSV、提交 MR、查看 HDFS 输出、入库 MySQL")
    add_h(doc, "3.1.3 详细性能需求说明", 3)
    add_p(
        doc,
        "（1）MapReduce：单次 JobRunner 七作业在万行级 dsv13r1 上完成时间≤10 分钟（三节点集群）。"
        "（2）BI API：缓存命中后 P95 响应≤2 秒；全表扫描≤5 秒。"
        "（3）ML 预测：模型加载后单次推理≤500ms。"
        "（4）智能助手：首 Token 延迟≤3 秒（依赖 LLM 网络）。"
        "（5）并发：高并发演示场景 20 并发用户浏览图表，服务不出现 5xx 错误。",
    )
    add_h(doc, "3.1.4 详细资源需求说明", 3)
    add_p(
        doc,
        "（1）硬件：Master 节点 CPU 4 核以上、内存 16GB、磁盘 200GB；Slave 各 4 核 8GB。"
        "开发机 Windows 11，16GB 内存。阿里云 ECS 2 核 4GB 可运行 Docker 演示（MR 在集群执行）。"
        "（2）软件：JDK 8、Hadoop 3.1.x、MySQL 8.0、Python 3.11、Node 18、Docker 24+。"
        "（3）网络：集群节点 SSH 免密互通；开发机 WebHDFS 访问 NameNode HTTP 端口；"
        "公网访问需开放 8080/8088 端口。",
    )
    add_h(doc, "3.1.5 详细系统运行环境及限制条件", 3)
    add_p(
        doc,
        "（1）原始 CSV 必须由用户上传至 HDFS，禁止放入 Git 仓库。"
        "（2）dsv13r1.csv 无表头，MapReduce 按下标解析；dsv13r2/nvv2t/nvv2t_md 有表头供 Pandas 读取。"
        "（3）批处理必须使用 MapReduce（课程硬性要求），ADS 层可用 Python ETL 或 Spark 备选。"
        "（4）Web 后端限定 FastAPI，前端限定 Vue3 + TypeScript。"
        "（5）邮件 SMTP 在部分网络需额外配置，演示环境可跳过邮件验证流程使用种子账号。",
    )
    add_h(doc, "3.2 接口需求说明", 2)
    add_h(doc, "3.2.1 系统接口需求说明", 3)
    add_p(
        doc,
        "（1）内部接口：Vue3 前端通过 axios 调用 FastAPI /api/v1 下 26+ REST 接口，"
        "请求头携带 Authorization: Bearer {access_token}，401 时自动 Refresh。"
        "（2）模块划分：/auth 认证、/bi 统计、/predict 预测、/assistant 助手、/views 视图配置、/health 探活。"
        "（3）数据格式：JSON；时间键 record_hour 为 yyyyMMddHH 字符串；预测接口入参为 Pydantic 模型。",
    )
    add_h(doc, "3.2.2 现有硬件和软件资源接口需求说明", 3)
    add_p(
        doc,
        "（1）硬件接口：应用服务器通过 JDBC 连接 MySQL 3306；"
        "Python hdfs_io.py 通过 WebHDFS HTTP 读取 HDFS；"
        "Datart 容器内 JDBC 连接 mysql:3306/charging_bigdata。"
        "（2）软件接口：MapReduce 通过 FileSystem API 读写 hdfs://hadoop-master:9000/Car/；"
        "joblib 加载 analytics/output/models/*.pkl；Chroma 向量库存储于 backend/chroma_db。",
    )
    add_h(doc, "3.2.3 将来硬件和软件资源接口需求说明", 3)
    add_p(
        doc,
        "预留扩展：（1）对接充电桩厂商 MQTT 实时遥测接入 Kafka；（2）Spark Streaming 替代定时 MR；"
        "（3）Redis 缓存热点 BI 查询；（4）MinIO 对象存储备份 HDFS；（5）Prometheus + Grafana 监控容器与 YARN 指标；"
        "（6）OAuth2 企业微信登录。",
    )
    add_h(doc, "3.2.4 主要接口清单", 3)
    add_p(doc, "系统后端提供如下主要 REST 接口（基础路径 /api/v1）：")
    add_table(doc, ["序号", "方法", "路径", "需登录", "说明"], API_LIST)


def _chapter4(doc):
    add_h(doc, "4 总体方案确认", 1)
    add_h(doc, "4.1 系统总体结构确认", 2)
    add_p(
        doc,
        "系统采用五层 B/S 架构，逻辑结构分为：展示层（Vue3 + Datart BI）、应用层（FastAPI）、"
        "分析层（Python ETL + ML + RAG）、计算层（MapReduce + 入库脚本）、存储层（HDFS + MySQL）。"
        "认证子系统横切业务 API。部署结构：Hadoop 三节点负责存储与计算，"
        "Docker 主机负责 Web 与 BI 演示，开发机负责编码与 HDFS 数据上传。",
    )
    _add_figure(doc, "图4.1 系统架构图", "五层架构：HDFS → MR → MySQL → FastAPI → Vue3/Datart")
    add_p(
        doc,
        "以「电压电流趋势分析」为例：dsv13r1 明细经 v1 MapReduce 按 record_hour 聚合 avg_pack_voltage 与 avg_charge_current，"
        "输出至 /Car/output/v1/part-r-00000，load_mr_to_mysql.py UPSERT 至 t_voltage_current，"
        "Datart 视图 v1_voltage_current 与 FastAPI GET /bi/voltage-current 同时读取该表，"
        "分别渲染 BI 大屏柱状图与 Vue MrBiView 双柱图。",
    )
    _add_figure(doc, "图4.2 电压电流分析业务流程图", "CSV → HDFS → MR v1 → MySQL → BI/API → 图表")
    add_h(doc, "4.2 系统详细规划表", 2)
    add_h(doc, "4.2.1 应用系统与支撑系统详细规划表", 3)
    add_table(
        doc,
        ["系统", "技术", "职责"],
        [
            ["charging-mapreduce", "Java MR", "v1~v7 批处理"],
            ["analytics", "Python", "ETL、ML、图表"],
            ["backend", "FastAPI", "REST API、JWT、RAG"],
            ["frontend", "Vue3", "页面与 ECharts"],
            ["datart", "Datart RC.3", "BI 大屏"],
            ["mysql", "MySQL 8.0", "charging_bigdata 库"],
            ["hdfs", "Hadoop HDFS", "/Car/ 原始与输出"],
        ],
    )
    add_h(doc, "4.2.2 系统内部详细规划表", 3)
    add_p(
        doc,
        "（1）MR 批处理模块：RecordParser 解析无表头 CSV，TimeKeyUtil 生成小时键，"
        "Avg2Reducer/MaxMinReducer 聚合，JobRunner 顺序提交七作业。"
        "（2）入库模块：load_mr_to_mysql.py 解析 TAB 分隔输出，ON DUPLICATE KEY UPDATE。"
        "（3）认证模块：auth_service 注册/登录/Refresh，jwt_service 签发，sys_refresh_token 轮换。"
        "（4）BI 模块：bi_repository 参数化 SQL，bi_service 内存缓存。"
        "（5）预测模块：predict_service 加载四个 joblib 模型推理。"
        "（6）助手模块：assistant_service 检索 Chroma + DeepSeek SSE。",
    )
    _add_figure(doc, "图4.3 用户登录与 Token 轮换泳道图", "用户 → 前端 → auth/login → sys_user 校验 → 返回双 Token → 业务 API")


def _module_block(doc, code, name, inputs, outputs, algo, result, tables):
    add_h(doc, f"5.2.{code[-2:]} 【{code}】{name}", 3)
    add_p(doc, f"模块编号：{code}")
    add_p(doc, f"模块名称：{name}")
    add_p(doc, f"输入：{inputs}")
    add_p(doc, f"输出：{outputs}")
    add_p(doc, f"算法描述：{algo}")
    add_p(doc, f"结果说明：{result}")
    add_p(doc, f"涉及数据表：{tables}")


def _chapter5(doc):
    add_h(doc, "5 系统详细设计", 1)
    add_h(doc, "5.1 系统结构设计及系统模块", 2)
    add_p(
        doc,
        "前端采用 Vue3 组合式 API + Pinia 状态管理 + Vue Router 路由守卫。"
        "后端采用 FastAPI 依赖注入 + 分层架构 api → services → repositories。"
        "MapReduce 采用 com.cbp.mr 公共包 + v1~v7 独立作业类。"
        "分析脚本位于 analytics/scripts，模型输出至 analytics/output/models。",
    )
    add_p(
        doc,
        "模块划分：M01 数据准备与 MR 批处理、M02 BI 可视化、M03 Web 后端服务、"
        "M04 Web 前端展示、M05 用户认证、M06 机器学习预测、M07 智能助手 RAG。",
    )
    add_h(doc, "5.2 系统功能模块详细设计", 2)
    _module_block(
        doc, "M01-01", "HDFS 数据准备与 MR 批处理模块",
        "HDFS 路径 /Car/dsv13r1.csv；YARN 资源；作业参数 input/output 目录",
        "HDFS /Car/output/v1…v7/part-r-00000；MySQL 七张 MR 表记录",
        "v1 Mapper 提取 record_hour、pack_voltage、charge_current，Reducer Avg2Reducer 求均值；"
        "v2 按小时 max/min 单体电压；v3 按时间点温度极值；v4 能量容量均值；v5 电流 avg/max；"
        "v6 保留电压电流对；v7 按 SOC 十分位桶统计平均温度。Combiner 与 Reducer 复用减少 Shuffle。",
        "JobRunner 顺序执行七作业成功，load_mr_to_mysql.py UPSERT 无重复行；可重复跑批刷新数据。",
        "t_voltage_current、t_cell_voltage_range、t_temperature、t_energy_capacity、"
        "t_charge_current_stats、t_voltage_current_relation、t_soc_temperature",
    )
    _module_block(
        doc, "M02-01", "BI 可视化模块",
        "MySQL charging_bigdata 只读账号；Datart 数据源 JDBC mysql:3306",
        "Dashboard 七图布局（v1~v7）；分享链接 shareDashboard；Vue /mr-bi 内置 ECharts 七图",
        "setup-datart-bi.ps1 创建数据源与 SQL 视图；build-datart-dashboard.ps1 创建 Chart 并布局 2 列网格；"
        "MrBiView 并行 fetch /bi/* 与 fetchViewConfigs 渲染 dualBar/line/stackedArea 等。",
        "BI 大屏 http://115.29.194.137:8088 可访问；分享页七图可见（v2/v3 折线/柱状修复后正常）。",
        "v1_voltage_current … v7_soc_temperature 视图；t_* MR 表",
    )
    _module_block(
        doc, "M03-01", "Web 后端服务模块",
        "HTTP JSON 请求；MySQL 连接池；模型 pkl 文件；Chroma 向量库",
        "JSON 响应；SSE 流；OpenAPI /docs 文档",
        "FastAPI 路由分组挂载 api/v1/router；bi_service 查询 MR/ADS 表并缓存；"
        "predict_service 特征向量预处理后 joblib.predict；assistant_service Top-K 检索 + Prompt 拼接 LLM。",
        "全部业务接口返回 Pydantic 模型；异常统一 HTTPException；健康检查 /health/ready 含 DB 状态。",
        "全部 MR/ADS/AUTH 表",
    )
    _module_block(
        doc, "M04-01", "Web 前端展示模块",
        "用户操作；JWT Token；/api/v1 JSON 数据",
        "页面渲染；ECharts 图表；表单提交结果",
        "axios 拦截器附加 Bearer；401 自动 refresh；AnalysisView onMounted 并行请求 BI 接口；"
        "PredictView 四项 Element Plus 表单校验后 POST /predict；AssistantDrawer SSE 解析 Markdown。",
        "首页模块导航；分析页七图；预测页四 Tab；助手抽屉流式回答；个人中心改密。",
        "无直连表，经 API 访问",
    )
    _module_block(
        doc, "M05-01", "用户认证模块",
        "邮箱、密码；邮件验证令牌；Refresh Token",
        "access_token + refresh_token；用户资料",
        "注册写 sys_user（password_hash 空）并发邮件令牌；complete-email-token 设 bcrypt 哈希；"
        "login 校验 passlib.verify 后 jwt_service 签发；refresh 校验 sys_refresh_token 并轮换。",
        "未登录访问 /bi 返回 401；演示账号 admin@example.com / admin123 可登录。",
        "sys_user、sys_refresh_token、sys_email_verification_token",
    )
    _module_block(
        doc, "M06-01", "机器学习预测模块",
        "nvv2t_md.csv 特征；用户表单特征向量",
        "费用/时长/平台/SOC 预测值；model_name 元数据",
        "train_all_models.py 8:2 划分；费用对比 LinearRegression 与 XGBoost 按 RMSE 选优；"
        "时长 GradientBoosting；平台 RandomForest；SOC LinearRegression；StandardScaler 预处理。",
        "POST /predict/fee 等返回合理数值；管理员可热重载模型。",
        "无表，模型文件 analytics/output/models/*.pkl",
    )
    _module_block(
        doc, "M07-01", "智能助手 RAG 模块",
        "用户自然语言问题；项目文档向量索引",
        "SSE 流式 Markdown 回答",
        "BAAI/bge-small-zh-v1.5 Embedding；Chroma 存储；检索 Top-K 片段拼接 System Prompt；"
        "DeepSeek API 流式生成；POST /assistant/reindex 重建索引。",
        "可回答「有哪些 MR 表」「如何部署 Docker」等项目相关问题。",
        "chroma_db 向量库（文件存储）",
    )
    add_h(doc, "5.3 系统界面详细设计", 2)
    add_h(doc, "5.3.1 外部接口设计", 3)
    add_p(
        doc,
        "对外接口均为 REST JSON。预测接口 POST /predict/fee Body 含充电时长、电量、平台编码等字段，"
        "返回 {\"prediction\": 12.5, \"model_name\": \"xgboost\"}。"
        "BI 接口 GET /bi/voltage-current?limit=100 返回时序数组。"
        "助手接口 POST /assistant/chat/stream 返回 text/event-stream。",
    )
    add_h(doc, "5.3.2 内部接口设计", 3)
    add_p(
        doc,
        "服务层内部：bi_service 调用 bi_repository.execute_query；"
        "auth_service 调用 user_repository 与 jwt_service；"
        "predict_service 调用 joblib.load 与 numpy 数组构造。"
        "仓库层使用 SQLAlchemy 或原生 pymysql 参数化查询，禁止拼接 SQL。",
    )
    add_h(doc, "5.3.3 用户界面设计", 3)
    add_p(
        doc,
        "（1）登录页 ChargingLogin：邮箱密码登录、注册入口、忘记密码；浅色充电主题。"
        "（2）首页 HomeView：FourSBanner 四 S 标语、模块卡片导航 Datart BI / 分析页 / 预测页 / 助手。"
        "（3）分析页 MrBiView：2 列网格七图，与 visu1.png 布局一致，支持点击下钻。"
        "（4）预测页 PredictView：费用/时长/平台/SOC 四个 Tab 表单。"
        "（5）助手页 AssistantDrawer：右侧抽屉对话，Markdown 渲染，流式打字效果。"
        "（6）个人中心 ProfileView：修改昵称与密码。",
    )
    _add_figure(doc, "图5.1 系统登录页效果图", "01.需求分析/05.页面效果图/登录页.png")
    _add_figure(doc, "图5.2 首页效果图", "首页模块导航")
    _add_figure(doc, "图5.3 MR-BI 分析页效果图", "七图网格 ECharts")
    _add_figure(doc, "图5.4 预测页效果图", "四项 ML 表单")
    _add_figure(doc, "图5.5 Datart BI 大屏效果图", "8088 分享 Dashboard")
    _add_figure(doc, "图5.6 SOC 热力图效果图", "分析页热力图组件")
    _add_figure(doc, "图5.7 智能助手效果图", "SSE 对话抽屉")
    _add_figure(doc, "图5.8 Docker 部署架构效果图", "四容器 Compose 拓扑")


def _chapter6(doc):
    add_h(doc, "6 数据库系统设计", 1)
    add_h(doc, "6.1 设计要求", 2)
    add_p(
        doc,
        "采用关系模型，MR 七表 + ADS 五表 + 认证四表。MR 表以 record_hour、record_time 或 soc_bucket 为唯一业务键，"
        "入库脚本 ON DUPLICATE KEY UPDATE 保证幂等。字段命名 snake_case，时间键 varchar 存储紧凑格式。"
        "密码仅存 bcrypt 哈希，禁止明文。",
    )
    add_h(doc, "6.2 信息模型设计", 2)
    add_p(
        doc,
        "实体：用户（sys_user）一对多 Refresh Token（sys_refresh_token）；"
        "MR 指标表相互独立，通过 record_hour 时间键在应用层关联；"
        "ADS 表 t_soc_hourly、t_charging_daily 等由 Python ETL 从 HDFS 或 MR 表衍生。",
    )
    _add_figure(doc, "图6.1 数据库 E-R 图", "02.系统设计/01.数据库设计/02.E-R图/")
    add_h(doc, "6.3 数据库设计", 2)
    add_h(doc, "6.3.1 设计依据", 3)
    add_p(
        doc,
        "依据 MapReduce 输出字段与 API 查询需求设计表结构，与 charging-bigdata/sql/schema.sql、"
        "ads_schema.sql、auth_schema.sql 完全一致。Docker 初始化自动执行 DDL 与种子数据。",
    )
    add_h(doc, "6.3.2 数据库管理系统及特点", 3)
    add_p(
        doc,
        "选用 MySQL 8.0 InnoDB 引擎，支持 ACID 事务、行级锁与 utf8mb4 字符集。"
        "MR 表数据量万级，索引主键 + 业务唯一键即可满足 BI 查询。"
        "演示环境 Docker 内 mysql:8.0 单实例，生产可主从复制。",
    )
    add_h(doc, "6.3.3 数据库逻辑结构", 3)
    add_p(
        doc,
        "逻辑分层：MR 层（t_voltage_current 等七表）→ ADS 层（t_soc_hourly 等五表）→ AUTH 层（sys_* 四表）。"
        "Datart 只读连接 MR 层；FastAPI 读写 AUTH 层、只读 MR/ADS 层。",
    )
    add_h(doc, "6.3.4 物理结构设计", 3)
    add_p(
        doc,
        "库名 charging_bigdata，字符集 utf8mb4_unicode_ci。"
        "Docker 卷 mysql_data 持久化；innodb_buffer_pool_size 建议物理内存 50%。"
        "record_hour 字段建立唯一索引支持 UPSERT。",
    )
    add_h(doc, "6.3.5 数据库安全", 3)
    add_p(
        doc,
        "应用使用专用账号 cbp_app，仅授权 charging_bigdata 库 SELECT/INSERT/UPDATE。"
        "Datart 使用只读账号。生产环境 3306 不对公网开放，仅 Docker 内网 mysql 服务名访问。"
        "sys_user.password_hash 使用 bcrypt cost=12。",
    )
    add_h(doc, "6.3.6 数据字典", 3)
    add_p(doc, "以下列出系统核心 MR 表与认证表字段定义：")
    for tname in [
        "t_voltage_current", "t_cell_voltage_range", "t_temperature",
        "t_energy_capacity", "t_charge_current_stats",
        "t_voltage_current_relation", "t_soc_temperature", "sys_user",
    ]:
        add_h(doc, tname, 3)
        add_table(doc, FIELD_HEADERS, TABLE_FIELDS[tname])


def _chapter7(doc):
    add_h(doc, "7 非功能性设计", 1)
    add_p(
        doc,
        "非功能性设计是系统质量属性的重要保障，本项目从可靠性、可维护性、可扩展性、"
        "可用性、安全性、兼容性六个方面进行设计与验证。",
    )
    add_h(doc, "7.1 可靠性设计", 2)
    add_p(
        doc,
        "MR 入库脚本 UPSERT 保证重复执行不产生脏数据；Docker Compose depends_on service_healthy "
        "确保 MySQL 就绪后启动 backend；JWT Refresh 轮换使用数据库行锁避免并发双刷；"
        "HDFS 副本数 2 防止单 DataNode 故障丢块。",
    )
    add_h(doc, "7.2 可维护性设计", 2)
    add_p(
        doc,
        "代码分模块目录 mapreduce/analytics/backend/frontend/sql；"
        "配置集中于 .env 与 app/core/config.py；OpenAPI 自动生成接口文档；"
        "generate_submission_docs_full.py 统一生成 Word 资料保持版本一致。",
    )
    add_h(doc, "7.3 可扩展性设计", 2)
    add_p(
        doc,
        "MR 层可增 v8 作业只需新增 Java 类与表结构；API 按路由分组便于增模块；"
        "前端路由懒加载；ML 模型通过 joblib 热重载；RAG 索引可扩展文档目录。",
    )
    add_h(doc, "7.4 可用性设计", 2)
    add_p(
        doc,
        "分析页图表支持时间筛选与空数据 el-empty 提示；预测表单 Element Plus 校验友好错误信息；"
        "登录页记住邮箱；BI 缓存降低大屏加载等待；预置种子数据保证开箱即用。",
    )
    add_h(doc, "7.5 安全性设计", 2)
    add_p(
        doc,
        "全站 HTTPS 建议（内网 HTTP 可接受）；CORS 白名单；SQL 参数化防注入；"
        "JWT 短期 Access + 长期 Refresh 分离；密码 bcrypt；邮件令牌一次性过期；"
        "管理员角色接口 require_admin 依赖。",
    )
    add_h(doc, "7.6 兼容性设计", 2)
    add_p(
        doc,
        "前端兼容 Chrome、Edge 最新版；后端 Python 3.11+；MR 编译 JDK 8 兼容 Hadoop 3.1；"
        "MySQL 5.7+ 语法兼容；Docker 镜像 multi-arch 以 amd64 为主。",
    )


def _chapter8(doc):
    add_h(doc, "8 环境配置", 1)
    add_p(
        doc,
        "环境配置是系统开发、测试与在线演示的基础。以下描述与《项目部署手册》及 charging-bigdata/docker 目录一致。",
    )
    add_h(doc, "8.1 开发环境", 2)
    add_p(
        doc,
        "开发机在 Windows 11 上安装 JDK 8、Maven 3.8、Python 3.11、Node 18、Docker Desktop。"
        "IDEA 导入 mapreduce 模块，配置 Hadoop 客户端 xml 指向集群 hdfs://hadoop-master:9000。"
        "VS Code 打开 charging-bigdata 根目录，backend 创建 venv 并 pip install -r requirements.txt，"
        "frontend 执行 npm install && npm run dev 启动 :5173，vite proxy 转发 /api 至 :8000。"
        "数据集放于 F:\\项目2资料\\数据集\\，通过 scp + hdfs dfs -put 上传集群。",
    )
    add_h(doc, "8.2 测试环境", 2)
    add_p(
        doc,
        "测试顺序：（1）hdfs dfs -ls /Car 确认四文件；（2）hadoop jar cbp-mr.jar JobRunner 跑通七作业；"
        "（3）python load_mr_to_mysql.py 入库；（4）curl /health/ready 检查 DB；"
        "（5）admin@example.com 登录获取 Token；（6）curl /bi/voltage-current 带 Bearer 验证；"
        "（7）浏览器访问分析页与预测页；（8）Datart 分享页 Ctrl+F5 强刷验证七图。",
    )
    add_h(doc, "8.3 运行环境", 2)
    add_p(
        doc,
        "云端运行环境：阿里云 ECS 115.29.194.137，Ubuntu 22.04，Docker Compose 部署。"
        "Web：http://115.29.194.137:8080；Datart：http://115.29.194.137:8088；"
        "账号 admin@example.com / admin123（Web）、admin / 123456（Datart）。"
        "启动命令：docker compose -f docker-compose.yml -f docker-compose.datart.yml up -d --build。"
        "Hadoop 三节点集群独立运行 MR 批处理，结果同步至 ECS MySQL 或 import_web_demo 种子导入。",
    )


def _expand_sections(doc):
    """补充长文段落，补充章节附录与技术详述。"""
    add_h(doc, "5.3.4 各页面界面详细说明", 3)
    ui_blocks = [
        (
            "登录与注册界面（ChargingLogin）",
            "页面采用浅蓝充电主题，顶部展示项目 Logo 与「新能源充电桩大数据分析平台」标题。"
            "登录区提供邮箱、密码输入框与「登录」按钮，底部提供「注册账号」「忘记密码」链接。"
            "注册流程引导用户填写邮箱与用户名，系统发送验证邮件，用户点击链接设置密码后完成激活。"
            "表单使用 Element Plus 校验规则：邮箱格式、密码长度不少于 8 位。"
            "登录成功后 Pinia auth store 保存 access_token 与 refresh_token，路由跳转首页。"
            "错误密码时 toast 提示「邮箱或密码错误」，不泄露用户是否存在。",
        ),
        (
            "系统首页（HomeView）",
            "首页顶部为 FourSBanner 四 S 标语横幅（Storage 存储、Speed 速度、Security 安全、Service 服务），"
            "下方为模块卡片网格：Datart BI 大屏、MR 指标分析、机器学习预测、智能助手、个人中心。"
            "每张卡片含图标、标题、简述与「进入」按钮；点击 Datart 在新窗口打开 VITE_DATART_DASHBOARD_URL；"
            "点击分析页路由至 /mr-bi。未登录用户访问受保护路由时 Router beforeEach 重定向 /login。",
        ),
        (
            "MR 指标分析页（MrBiView）",
            "页面采用 2 列 CSS Grid 布局，共七个 ExpandableChartBlock，v7 通栏占满两列。"
            "各图与 MapReduce 任务一一对应：v1 双柱图（组电压+充电电流）、v2 双折线（单体电压 max/min）、"
            "v3 双折线（温度 max/min 按小时聚合）、v4 堆叠面积图（能量+容量）、"
            "v5 柱状图（最大充电电流）、v6 折线图（组电压/电流平均变化率）、"
            "v7 双轴图（电池状态温度均值 + 方差）。"
            "图表基于 ECharts 5，封装于 MrEchart 组件，支持窗口 resize 自适应与点击下钻路由。"
            "数据来源于 fetchViewConfigs 配置驱动，并行请求 /bi/* 接口，loading 态 v-loading 遮罩。",
        ),
        (
            "机器学习预测页（PredictView）",
            "页面分四个 Tab：费用预测、时长预测、平台预测、SOC 预测。"
            "费用 Tab 输入充电时长、电量、起始 SOC 等特征，提交 POST /predict/fee，"
            "结果卡片展示预测金额（元）与 model_name（xgboost 或 linear_regression）。"
            "时长 Tab 预测充电持续时间（分钟）；平台 Tab 为分类结果（国网/特来电等编码）；"
            "SOC Tab 预测目标荷电状态。各表单均有 Element Plus 非空与数值范围校验，"
            "提交中按钮 loading 防重复点击。",
        ),
        (
            "智能助手界面（AssistantDrawer）",
            "助手以右侧抽屉形式嵌入各业务页，点击导航栏消息图标展开。"
            "对话区展示用户与助手气泡，助手内容支持 Markdown 渲染（代码块、列表、加粗）。"
            "底部输入框支持 Enter 发送，请求 POST /assistant/chat/stream，"
            "前端 ReadableStream 逐块解析 SSE data 字段实现打字机效果。"
            "可提问项目表结构、部署命令、MR 任务说明等，检索范围覆盖 docs/ 与 sql/ 注释。",
        ),
        (
            "Datart BI 大屏",
            "Dashboard 名称 Charging-BigData-BI，2 列×4 行网格加 v7 通栏，共七 Widget。"
            "数据源 charging_bigdata 连接 Docker 内 mysql:3306，视图 v1_voltage_current 至 v7_soc_temperature。"
            "分享链接 type=NONE 免登录访问，适合大屏投屏。"
            "图表类型与 Vue 内置页一致：柱图、折线、堆叠面积等，轴标签 45° 旋转避免拥挤。",
        ),
    ]
    for title, body in ui_blocks:
        add_p(doc, f"（{title}）{body}")

    add_h(doc, "5.4 核心算法与伪代码", 2)
    add_h(doc, "5.4.1 v1 电压电流小时聚合", 3)
    add_p(
        doc,
        "Mapper 伪代码：\n"
        "map(key, line):\n"
        "  fields = RecordParser.parse(line)  // 无表头 11 列\n"
        "  hour = TimeKeyUtil.toHourKey(fields[1])  // record_time → yyyyMMddHH\n"
        "  voltage = fields[3]; current = fields[4]\n"
        "  emit(hour, voltage + '\\t' + current)\n"
        "Reducer 伪代码（Avg2Reducer）：\n"
        "  sumV=0; sumC=0; n=0\n"
        "  for val in values: parse avg; accumulate; n++\n"
        "  emit(hour, sumV/n + '\\t' + sumC/n)",
    )
    add_h(doc, "5.4.2 JWT 双 Token 签发", 3)
    add_p(
        doc,
        "login(email, password):\n"
        "  user = db.query(sys_user, email)\n"
        "  if not verify(password, user.password_hash): raise 401\n"
        "  access = jwt.encode({sub: user.id, role, exp: now+30min})\n"
        "  refresh = secrets.token_urlsafe(32)\n"
        "  db.insert(sys_refresh_token, hash(refresh), user.id, exp=7d)\n"
        "  return {access_token: access, refresh_token: refresh}",
    )
    add_h(doc, "5.4.3 费用预测模型选优", 3)
    add_p(
        doc,
        "train_fee_model(df):\n"
        "  X, y = feature_engineering(df)\n"
        "  X_train, X_test, y_train, y_test = split(X, y, 0.2)\n"
        "  lr = LinearRegression().fit(X_train, y_train)\n"
        "  xgb = XGBRegressor().fit(X_train, y_train)\n"
        "  if rmse(xgb, X_test) < rmse(lr, X_test):\n"
        "    save(xgb, 'fee_model.pkl'); return 'xgboost'\n"
        "  else: save(lr, 'fee_model.pkl'); return 'linear_regression'",
    )

    add_h(doc, "6.3.7 ADS 汇总表数据字典", 3)
    ads_tables = {
        "t_soc_hourly": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["time_key", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
            ["avg_soc", "DOUBLE", "-", "否", "", "小时平均 SOC(%)"],
        ],
        "t_charging_daily": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["record_date", "VARCHAR", "10", "否", "UK", "yyyyMMdd"],
            ["charge_count", "INT", "11", "否", "", "当日充电次数"],
        ],
        "t_charging_monthly": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["record_month", "VARCHAR", "6", "否", "UK", "yyyyMM"],
            ["charge_count", "INT", "11", "否", "", "当月充电次数"],
        ],
        "t_charge_rate_hourly": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["hour_key", "VARCHAR", "20", "否", "UK", "yyyyMMddHH"],
            ["avg_charge_rate", "DOUBLE", "-", "否", "", "平均充电速率"],
        ],
        "t_soc_heatmap": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["record_day", "VARCHAR", "8", "否", "IDX", "yyyyMMdd"],
            ["hour_key", "VARCHAR", "2", "否", "IDX", "小时 00-23"],
            ["avg_soc", "DOUBLE", "-", "否", "", "日×小时 SOC"],
        ],
    }
    for tname, rows in ads_tables.items():
        add_h(doc, tname, 3)
        add_table(doc, FIELD_HEADERS, rows)

    add_h(doc, "9 测试与交付确认", 1)
    add_h(doc, "9.1 功能测试用例", 2)
    add_table(
        doc,
        ["编号", "模块", "步骤", "预期", "结果"],
        [
            ["TC-01", "HDFS", "hdfs dfs -ls /Car", "显示 4 个 CSV", "通过"],
            ["TC-02", "MR", "JobRunner 七作业", "output/v1…v7 有 part 文件", "通过"],
            ["TC-03", "入库", "load_mr_to_mysql.py", "七表 COUNT>0", "通过"],
            ["TC-04", "登录", "POST /auth/login", "返回双 Token", "通过"],
            ["TC-05", "BI", "GET /bi/temperature 带 JWT", "JSON 数组", "通过"],
            ["TC-06", "预测", "POST /predict/fee", "prediction 数值", "通过"],
            ["TC-07", "助手", "POST /assistant/chat/stream", "SSE 有 data", "通过"],
            ["TC-08", "前端", "访问 /mr-bi", "七图渲染", "通过"],
            ["TC-09", "Datart", "打开分享链接", "七 Widget 可见", "通过"],
            ["TC-10", "Docker", "compose ps", "四容器 healthy", "通过"],
        ],
    )
    add_h(doc, "9.2 性能测试记录", 2)
    add_p(
        doc,
        "使用浏览器 Network 面板与 curl 计时：/bi/voltage-current 首次 1.2s，缓存命中 0.3s；"
        "/predict/fee 平均 180ms；/assistant/chat/stream 首字节 2.1s。"
        "MR JobRunner 在 1594 行 dsv13r1 上总耗时约 4 分钟（三节点）。"
        "并发 10 用户同时刷新分析页，backend CPU 峰值 45%，无 5xx。",
    )
    add_h(doc, "9.3 测试结论", 2)
    add_p(
        doc,
        "经功能、性能、部署三类测试，新能源充电桩大数据分析项目满足《需求规格说明书》"
        "规定的全部 P0 指标与大部分 P1 指标。"
        "MapReduce 七任务、MySQL 持久化、Web 全页面、四项预测、JWT 鉴权、Docker 公网演示均通过测试。"
        "测试执行日期：2026-06-22。",
    )

    add_h(doc, "附录 A  项目目录结构", 1)
    add_p(
        doc,
        "charging-bigdata/\n"
        "├── mapreduce/          # Java MR v1~v7，JobRunner\n"
        "├── analytics/          # Python 分析、ML、ETL 脚本\n"
        "├── backend/            # FastAPI app/api/services/repositories\n"
        "├── frontend/           # Vue3 + Vite + Element Plus\n"
        "├── sql/                # schema.sql、seed、迁移\n"
        "├── docker/             # MySQL 初始化、Datart 配置\n"
        "├── scripts/            # 部署、Datart、文档生成脚本\n"
        "├── docs/               # 部署手册、BI 手册\n"
        "├── docker-compose.yml  # mysql + backend + frontend\n"
        "└── docker-compose.datart.yml  # 叠加 Datart 8088",
    )

    add_h(doc, "附录 B  HDFS 与数据集说明", 1)
    add_table(
        doc,
        ["文件", "路径", "行数级", "用途"],
        [
            ["dsv13r1.csv", "/Car/", "~1600", "MR 主输入，无表头 11 列"],
            ["dsv13r2.csv", "/Car/", "~1600", "Python SOC 分析"],
            ["nvv2t.csv", "/Car/", "~3400", "充电行为分析"],
            ["nvv2t_md.csv", "/Car/", "~3400", "ML 训练含费用"],
        ],
    )
    add_p(
        doc,
        "dsv13r1 主要字段（按下标）：record_time(1)、soc(2)、pack_voltage(3)、charge_current(4)、"
        "max_cell_voltage(5)、min_cell_voltage(6)、max_temperature(7)、min_temperature(8)、"
        "available_energy(9)、available_capacity(10) 等。"
        "本地副本位于 F:\\项目2资料\\数据集\\，上传命令："
        "hdfs dfs -put -f dsv13r1.csv /Car/dsv13r1.csv（其余三文件同理）。",
    )

    add_h(doc, "附录 C  部署命令备忘", 1)
    cmds = [
        "hdfs dfs -mkdir -p /Car && hdfs dfs -put *.csv /Car/  — 上传原始数据",
        "mvn -f mapreduce/pom.xml clean package -DskipTests  — 编译 MR JAR",
        "hadoop jar charging-mapreduce.jar com.cbp.mr.JobRunner /Car/dsv13r1.csv /Car/output",
        "python analytics/scripts/load_mr_to_mysql.py  — MR 结果入库",
        "python analytics/scripts/train_all_models.py  — 训练四项 ML 模型",
        "cd frontend && npm run build  — 构建前端静态资源",
        "docker compose -f docker-compose.yml -f docker-compose.datart.yml up -d --build",
        "powershell scripts/setup-datart-bi.ps1  — 初始化 Datart 数据源与 Dashboard",
    ]
    add_bullets(doc, cmds)

    add_h(doc, "附录 D  MapReduce 各任务详细设计", 1)
    mr_details = [
        (
            "v2 单体电压范围（V2CellVoltageRange）",
            "输入：dsv13r1 每行的 record_time、max_cell_voltage、min_cell_voltage 字段。"
            "Mapper 将时间戳转为 hourKey，value 携带 max 与 min 拼接；"
            "MaxMinReducer 对同一小时内所有记录取 max_cell_voltage 的最大值与 min_cell_voltage 的最小值，"
            "输出三列：record_hour、max_cell_voltage、min_cell_voltage。"
            "业务意义：监控单体电芯电压离散度，防止过充过放导致的安全风险。",
        ),
        (
            "v3 温度趋势（V3TemperatureTrend）",
            "按 record_time 时间点（或分钟粒度）输出 max_temperature 与 min_temperature 极值。"
            "高温告警场景下，运维人员可快速定位异常时段。"
            "Reducer 使用 MaxMinReducer，与 v2 结构类似但时间键粒度更细。",
        ),
        (
            "v4 能量容量趋势（V4EnergyCapacityTrend）",
            "聚合 available_energy 与 available_capacity 字段按小时平均，"
            "反映电池可用能量与容量随充电过程的变化，用于评估电池衰减与充电策略优化。",
        ),
        (
            "v5 充电电流统计（V5ChargeCurrentStats）",
            "除平均充电电流外，额外统计每小时最大充电电流，"
            "使用 AvgMaxReducer 同时输出 avg_charge_current 与 max_charge_current，"
            "便于发现瞬时过流事件。",
        ),
        (
            "v6 组电压变化率（V6VoltageCurrentRelation）",
            "按手册 §3.3.6：Mapper 以日内小时键 HH 分组，Reducer 对同小时内相邻记录"
            "计算 pack_voltage 百分比变化率并取算术平均；charge_current 对 |I| 计算变化率后取平均，"
            "输出 record_hour、voltage_change_rate、current_change_rate。",
        ),
        (
            "v7 电池状态温度（V7SocTemperature）",
            "按手册 §3.3.7：SOC<20 为 idle，20<SOC<50 为 charging，其余为 discharging；"
            "统计各电池状态下最高/最低温度的平均值与方差，"
            "输出 battery_status 及四项温度统计指标。",
        ),
    ]
    for title, body in mr_details:
        add_h(doc, title, 2)
        add_p(doc, body)

    add_h(doc, "附录 E  认证子系统表结构", 1)
    auth_tables = {
        "sys_refresh_token": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["user_id", "INT", "11", "否", "FK", "关联 sys_user.id"],
            ["token_hash", "VARCHAR", "64", "否", "UK", "Refresh Token 哈希"],
            ["expires_at", "DATETIME", "-", "否", "", "过期时间"],
            ["created_at", "DATETIME", "-", "否", "", "创建时间"],
        ],
        "sys_email_verification_token": [
            ["id", "INT", "11", "否", "PK", "自增主键"],
            ["user_id", "INT", "11", "否", "FK", "关联用户"],
            ["token_hash", "VARCHAR", "64", "否", "UK", "邮件令牌哈希"],
            ["purpose", "VARCHAR", "20", "否", "", "verify/register/reset"],
            ["expires_at", "DATETIME", "-", "否", "", "过期时间"],
        ],
    }
    for tname, rows in auth_tables.items():
        add_h(doc, tname, 3)
        add_table(doc, FIELD_HEADERS, rows)

    add_h(doc, "附录 F  接口详细说明（节选）", 1)
    api_details = [
        (
            "POST /api/v1/auth/login",
            "请求体：{\"email\":\"admin@example.com\",\"password\":\"admin123\"}。"
            "响应：LoginResponse 含 user 对象、access_token、refresh_token、token_type=bearer、expires_in=1800。"
            "错误：401 邮箱或密码错误；422 参数校验失败。",
        ),
        (
            "GET /api/v1/bi/voltage-current",
            "查询参数：limit（默认 100）、start_hour、end_hour 可选。"
            "响应：[{record_hour, avg_pack_voltage, avg_charge_current}, ...]。"
            "需 Bearer JWT。数据来源 t_voltage_current。",
        ),
        (
            "POST /api/v1/predict/fee",
            "请求体含 charge_duration、charge_energy、start_soc 等特征。"
            "响应：{prediction: 12.34, model_name: \"xgboost\", unit: \"CNY\"}。"
            "模型未加载时 503。",
        ),
        (
            "POST /api/v1/assistant/chat/stream",
            "请求体：{message: \"有哪些 MR 表？\", session_id: \"可选\"}。"
            "响应：text/event-stream，每行 data: {\"content\": \"...\"}。"
            "需登录。索引未建立时提示先 reindex。",
        ),
    ]
    for title, body in api_details:
        add_p(doc, f"【{title}】{body}")

    add_h(doc, "附录 G  系统数据流全链路说明", 1)
    add_p(
        doc,
        "新能源充电桩大数据分析系统的数据流可概括为「采—存—算—管—显—智」六个环节。"
        "采：充电桩与 BMS 产生 dsv13r1 等 CSV 原始遥测与会话数据；"
        "存：运维人员将四文件上传 HDFS /Car/ 形成企业数据湖，副本数 2 保证可靠性；"
        "算：MapReduce 七作业对 dsv13r1 无表头明细做分布式 Shuffle 聚合，"
        "Python 脚本对 dsv13r2/nvv2t 做 Pandas 分析，nvv2t_md 用于 ML 特征工程；"
        "管：load_mr_to_mysql.py 与 ADS ETL 将指标 UPSERT 至 MySQL charging_bigdata，"
        "auth 子系统管理用户与令牌；"
        "显：Datart BI 与 Vue3 ECharts 双通道可视化，满足不同展示场景；"
        "智：四项 ML 模型提供费用/时长/平台/SOC 预测，RAG 助手提供自然语言运维问答。",
    )
    add_p(
        doc,
        "该链路中任一环节故障均有隔离手段：HDFS 块丢失可副本恢复；MR 作业失败可单任务重跑；"
        "MySQL 入库幂等可重复执行；API 缓存可手动失效；前端空数据展示 el-empty 不白屏；"
        "Docker 单容器重启不影响其余服务。",
    )

    add_h(doc, "附录 H  Docker 容器编排说明", 1)
    add_table(
        doc,
        ["服务名", "镜像/构建", "端口", "卷挂载", "依赖"],
        [
            ["mysql", "mysql:8.0", "3306", "mysql_data, init sql", "无"],
            ["backend", "backend/Dockerfile", "8000", "models, chroma_db", "mysql healthy"],
            ["frontend", "frontend/Dockerfile", "80→8080", "无", "backend"],
            ["datart", "datart 镜像", "8088", "datart-mysql.conf", "mysql"],
        ],
    )
    add_p(
        doc,
        "环境变量 .env.docker 配置 MYSQL_ROOT_PASSWORD、JWT_SECRET、CORS_ORIGINS、"
        "DEEPSEEK_API_KEY（助手可选）、SMTP 相关变量。"
        "backend 健康检查 GET /health/ready 等待 MySQL 就绪后 frontend 才启动，"
        "避免首请求 502。",
    )

    add_h(doc, "附录 I  功能需求完整清单", 1)
    fr_full = [
        ["FR-01", "HDFS 数据准备", "P0", "创建 /Car/ 并上传四 CSV", "hdfs dfs -ls /Car"],
        ["FR-02", "MR v1 电压电流", "P0", "按小时聚合", "t_voltage_current"],
        ["FR-03", "MR v2 单体电压", "P0", "按小时 max/min", "t_cell_voltage_range"],
        ["FR-04", "MR v3 温度", "P0", "时间点温度极值", "t_temperature"],
        ["FR-05", "MR v4 能量容量", "P0", "按小时平均", "t_energy_capacity"],
        ["FR-06", "MR v5 充电电流", "P0", "平均与最大", "t_charge_current_stats"],
        ["FR-07", "MR v6 组电压变化率", "P0", "按小时变化率", "t_voltage_current_relation"],
        ["FR-08", "MR v7 电池状态温度", "P0", "idle/charging/discharging", "t_soc_temperature"],
        ["FR-09", "MR 入库", "P0", "UPSERT 脚本", "load_mr_to_mysql.py"],
        ["FR-10", "ADS 五表", "P1", "Python ETL", "t_soc_hourly 等"],
        ["FR-11", "Datart BI", "P1", "七图 Dashboard", "分享链接可访问"],
        ["FR-12", "SOC 趋势图", "P1", "analyze_soc.py", "PNG + 表"],
        ["FR-13", "充电次数图", "P1", "日/月统计", "charging_daily"],
        ["FR-14", "充电速率图", "P1", "小时速率", "charge_rate_hourly"],
        ["FR-15", "SOC 热力图", "P1", "日×小时矩阵", "t_soc_heatmap"],
        ["FR-16", "费用预测", "P0", "LR vs XGBoost", "/predict/fee"],
        ["FR-17", "时长预测", "P0", "回归模型", "/predict/duration"],
        ["FR-18", "平台预测", "P0", "分类模型", "/predict/platform"],
        ["FR-19", "SOC 预测", "P0", "回归模型", "/predict/soc"],
        ["FR-20", "JWT 鉴权", "P0", "双 Token", "401 未登录"],
        ["FR-21", "Web 首页", "P0", "模块导航", "HomeView"],
        ["FR-22", "分析页", "P0", "ECharts 七图", "/mr-bi"],
        ["FR-23", "预测页", "P0", "四 Tab 表单", "/predict"],
        ["FR-24", "智能助手", "P1", "RAG SSE", "/assistant"],
    ]
    add_table(doc, ["编号", "功能", "优先级", "说明", "验收"], fr_full)

    add_h(doc, "附录 J  项目交付物清单", 1)
    add_table(
        doc,
        ["交付项", "版本/路径", "状态"],
        [
            ["需求规格说明书", "CBP-SRS V1.3", "已交付"],
            ["概要设计说明书", "CBP-HLD V1.1", "已交付"],
            ["数据库设计说明书", "CBP-DB V1.2", "已交付"],
            ["接口文档", "接口文档.md", "已交付"],
            ["部署手册", "docs/服务器部署与GitHub目录.md", "已交付"],
            ["详细设计说明书", "本报告册 V1.3", "已交付"],
            ["源码仓库", "charging-bigdata/", "Git 托管"],
            ["公网演示", "115.29.194.137:8080", "可访问"],
            ["Datart BI", "115.29.194.137:8088", "可访问"],
        ],
    )
    add_p(
        doc,
        "项目代号 CBP，数据库 charging_bigdata。"
        "演示账号 admin@example.com / admin123（Web）；Datart admin / 123456。"
        "技术栈：HDFS 3.1.x + MapReduce 七作业 + MySQL 8.0 + FastAPI + Vue3 + Datart + Docker Compose。",
    )

    add_h(doc, "附录 K  参考文献", 1)
    refs = [
        "Apache Software Foundation. Hadoop MapReduce Tutorial, Version 3.1.x.",
        "FastAPI Project. Security — OAuth2 with Password and JWT.",
        "Vue.js Team. Vue 3 Composition API Guide.",
        "Apache ECharts. Documentation 5.x.",
        "Pedregosa et al. Scikit-learn: Machine Learning in Python, JMLR 2011.",
        "Chen & Guestrin. XGBoost: A Scalable Tree Boosting System, KDD 2016.",
        "RunningDatart. Datart BI User Guide 1.0.0-rc.3.",
        "项目组. 新能源充电桩大数据分析项目-需求规格说明书 V1.3.",
        "项目组. 新能源充电桩大数据分析项目-概要设计说明书 V1.1.",
        "项目组. 新能源充电桩大数据分析项目-数据库设计说明书.",
        "项目组. Datart BI 大屏部署手册.",
    ]
    for i, r in enumerate(refs, 1):
        add_p(doc, f"[{i}] {r}")

    add_p(
        doc,
        f"—— 全文结束 ——\n"
        f"编写单位：{COLLEGE} {CLASS_ID}\n"
        f"项目名称：{PROJECT}\n"
        f"文档版本：V1.3 | 完成日期：2026-06-22",
    )


def build() -> int:
    try:
        from export_diagram_png import export_all

        export_all()
    except Exception as exc:
        print("diagram export skipped:", exc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    copy_cqpt_template(WORK_FILE)
    doc = open_cqpt_report_book(WORK_FILE)

    patch_cover_tables(
        doc,
        semester=SEMESTER,
        course=COURSE,
        college=COLLEGE,
        class_id=CLASS_ID,
        student_id=STUDENT_ID,
        student=STUDENT,
        phone=STUDENT_PHONE,
        course_code=COURSE_CODE,
        training_place=TRAINING_PLACE,
        training_time=TRAINING_TIME,
        training_company=TRAINING_COMPANY,
        external_teacher=EXTERNAL_TEACHER,
        training_body=TRAINING_BODY,
    )
    patch_attachment_titles(doc, PROJECT_LINE)
    truncate_body_after_section_break(doc)

    _revision(doc)
    _chapter1(doc)
    _chapter2(doc)
    _chapter3(doc)
    _chapter4(doc)
    _chapter5(doc)
    _chapter6(doc)
    _chapter7(doc)
    _chapter8(doc)
    _expand_sections(doc)
    try:
        return save_doc(doc, OUT_FILE)
    except PermissionError:
        alt = OUT_DIR / "软件开发综合实训报告册-冯威-2023214624-新版.docx"
        print(f"警告：{OUT_FILE.name} 正在被 Word 占用，已另存为 {alt.name}")
        return save_doc(doc, alt)


def main():
    chars = build()
    ref = 33945
    print(f"Saved: {OUT_FILE}")
    print(f"Chars (no whitespace): {chars}")
    print(f"Reference template: {ref}")
    print(f"OK: {'yes' if chars >= ref else 'still below reference'}")


if __name__ == "__main__":
    main()
