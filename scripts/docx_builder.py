# -*- coding: utf-8 -*-
"""Word 文档构建工具（对齐公众监督模板：宋体正文小四 + 黑体标题 + Heading 样式）。"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_TITLE = "华文中宋"

# 正文：小四 12pt；标题：一级三号16pt、二级四号14pt、三级小四12pt加粗
SIZE_BODY = 12
SIZE_H1 = 16
SIZE_H2 = 14
SIZE_H3 = 12
SIZE_COVER = 24

SCRIPTS_DIR = Path(__file__).resolve().parent
ASSETS_DIR = SCRIPTS_DIR / "assets"
CQPT_LOGO = ASSETS_DIR / "cqpt_logo.png"
CQPT_TEMPLATE = Path(
    r"F:\公众监督系统\项目资料提交模板\05.软件开发综合实训报告册"
    r"\软件开发综合实训报告册-冯威-2023214624.docx"
)


def _set_rfonts(rfonts, font_name: str):
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_run_font(run, size=SIZE_BODY, bold=False, font_name=FONT_BODY):
    run.font.name = font_name
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    _set_rfonts(rfonts, font_name)


def _set_paragraph_format(p, align=None, line_spacing=1.5, space_before=0, space_after=0, first_line_indent=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def _configure_styles(doc: Document):
    """配置 Normal 与 Heading 1~3，与模板排版一致。"""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(SIZE_BODY)
    _set_rfonts(normal._element.rPr.rFonts, FONT_BODY)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5

    for level, size, before, after in [
        (1, SIZE_H1, 12, 6),
        (2, SIZE_H2, 6, 3),
        (3, SIZE_H3, 3, 3),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_HEADING
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        _set_rfonts(style._element.rPr.rFonts, FONT_HEADING)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    _configure_styles(doc)
    return doc


def _clear_document_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag != "sectPr":
            body.remove(child)


def copy_cqpt_template(out_path: Path) -> Path:
    """复制学校模板到输出路径，完整保留水印、样式与节属性。"""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if not CQPT_TEMPLATE.exists():
        raise FileNotFoundError(f"模板不存在: {CQPT_TEMPLATE}")
    shutil.copy2(CQPT_TEMPLATE, out_path)
    return out_path


def open_cqpt_report_book(path: Path) -> Document:
    """打开已复制的报告册模板。"""
    return Document(str(path))


def _set_cell_text(cell, text: str, *, replace_all: bool = False):
    """改单元格文字；replace_all=True 时清空单元格内全部段落（用于合并大单元格）。"""
    if replace_all:
        cell.text = text
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def patch_cover_tables(
    doc: Document,
    *,
    semester: str,
    course: str,
    college: str,
    class_id: str,
    student_id: str,
    student: str,
    phone: str,
    course_code: str,
    training_place: str,
    training_time: str,
    training_company: str,
    external_teacher: str,
    training_body: str,
):
    """更新封面学生信息表与实训正文表（不改版式）。"""
    if len(doc.tables) < 2:
        return
    t0 = doc.tables[0]
    cover_vals = [semester, course, college, class_id, student_id, student, phone]
    for ri, val in enumerate(cover_vals):
        if ri < len(t0.rows):
            _set_cell_text(t0.rows[ri].cells[1], val)

    t1 = doc.tables[1]
    if len(t1.rows) >= 4:
        _set_cell_text(t1.rows[0].cells[1], course)
        _set_cell_text(t1.rows[0].cells[3], course_code)
        _set_cell_text(t1.rows[1].cells[1], training_place)
        _set_cell_text(t1.rows[1].cells[3], training_time)
        _set_cell_text(t1.rows[2].cells[1], training_company)
        _set_cell_text(t1.rows[2].cells[3], external_teacher)
    if len(t1.rows) >= 6:
        for cell in t1.rows[5].cells:
            _set_cell_text(cell, training_body, replace_all=True)


def patch_attachment_titles(doc: Document, project_line: str):
    """更新附件封面项目名称（保留模板扉页样式）。"""
    candidates = []
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        if "文档标题" not in st:
            continue
        if "附件：" in para.text or "每小组" in para.text:
            continue
        if "详细设计" in para.text and len(para.text.strip()) < 20:
            continue
        if not para.text.strip():
            continue
        candidates.append(para)
    if not candidates:
        return
    para = candidates[0]
    if para.runs:
        para.runs[0].text = project_line
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = project_line


def truncate_body_after_section_break(doc: Document):
    """保留封面与附件扉页，删除分节符之后旧正文，准备写入 CBP 章节。"""
    body = doc.element.body
    children = list(body)
    break_at = None
    for i, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            break_at = i
    if break_at is None:
        return
    for child in children[break_at + 1 :]:
        if child.tag.split("}")[-1] != "sectPr":
            body.remove(child)


def new_report_book_doc() -> Document:
    """已弃用：请使用 copy_cqpt_template + open_cqpt_report_book。"""
    if not CQPT_TEMPLATE.exists():
        return new_doc()
    doc = Document(str(CQPT_TEMPLATE))
    _clear_document_body(doc)
    _configure_styles(doc)
    return doc


def _clone_header_footer_part(dst_part, src_part):
    dst_part.is_linked_to_previous = False
    dst_el = dst_part._element
    src_el = src_part._element
    for child in list(dst_el):
        dst_el.remove(child)
    for child in src_el:
        dst_el.append(deepcopy(child))


def restore_cqpt_headers_footers(doc: Document):
    """从学校模板恢复页眉水印与页脚页码。"""
    if not CQPT_TEMPLATE.exists():
        return
    tpl = Document(str(CQPT_TEMPLATE))
    _clone_header_footer_part(doc.sections[0].header, tpl.sections[0].header)
    _clone_header_footer_part(doc.sections[0].footer, tpl.sections[0].footer)
    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True


def add_content_section_break(doc: Document):
    """附件正文起新节，页眉水印与模板第二节一致。"""
    doc.add_section(WD_SECTION.NEW_PAGE)
    if len(doc.sections) >= 2:
        doc.sections[-1].header.is_linked_to_previous = True
        doc.sections[-1].footer.is_linked_to_previous = True


def add_center_title(doc: Document, text: str, size=SIZE_COVER, bold=True, font_name=FONT_HEADING):
    p = doc.add_paragraph()
    _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 1.5, 6, 6)
    r = p.add_run(text)
    set_run_font(r, size, bold, font_name)


def add_h(doc: Document, text: str, level: int = 1):
    """使用 Word Heading 1/2/3 样式（黑体），与公众监督模板一致。"""
    level = max(1, min(3, level))
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    set_run_font(run, sizes[level], True, FONT_HEADING)


def add_p(doc: Document, text: str, bold=False, align=None, first_line_indent=2):
    p = doc.add_paragraph(style="Normal")
    indent = Cm(0.74) if first_line_indent and align != WD_ALIGN_PARAGRAPH.CENTER else None
    _set_paragraph_format(p, align, 1.5, 0, 0, indent)
    r = p.add_run(text)
    set_run_font(r, SIZE_BODY, bold, FONT_BODY)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_format(p, None, 1.5)
        if p.runs:
            p.clear()
        r = p.add_run(item)
        set_run_font(r, SIZE_BODY, False, FONT_BODY)


def add_kv_table(doc: Document, rows: list[list[str]], col_widths=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for ri, (label, value) in enumerate(rows):
        c0 = table.rows[ri].cells[0]
        c1 = table.rows[ri].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, SIZE_BODY, True, FONT_BODY)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, SIZE_BODY, False, FONT_BODY)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    if not rows:
        rows = [[]]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, 11, True, FONT_HEADING)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run_font(r, 11, False, FONT_BODY)
    doc.add_paragraph()


def add_page_break(doc: Document):
    doc.add_page_break()


def add_table_of_contents(doc: Document):
    """插入 Word 目录域（与模板一致：标题「目录」+ TOC \\o \"1-3\"）。"""
    title_style = "缺省文本"
    for name in ("缺省文本", "Normal", "Heading 1"):
        if name in [s.name for s in doc.styles]:
            title_style = name
            break
    p_title = doc.add_paragraph("目录", style=title_style)
    _set_paragraph_format(p_title, WD_ALIGN_PARAGRAPH.CENTER, 1.5, 6, 6)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    doc.add_paragraph()


def add_cqpt_cover_header(doc: Document):
    """封面：校徽书法图 + Title 样式主标题 + 教务处落款。"""
    doc.add_paragraph()
    if CQPT_LOGO.exists():
        p_logo = doc.add_paragraph()
        _set_paragraph_format(p_logo, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
        p_logo.add_run().add_picture(str(CQPT_LOGO), width=Cm(10))
    else:
        for _ in range(2):
            doc.add_paragraph()
    doc.add_paragraph()
    p_title = doc.add_paragraph(style="Title")
    if p_title.runs:
        run = p_title.runs[0]
        run.text = "学生《软件开发综合实训》报告册"
    else:
        run = p_title.add_run("学生《软件开发综合实训》报告册")
    set_run_font(run, SIZE_COVER, False, FONT_TITLE)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
    r = p.add_run("重庆邮电大学教务处制")
    set_run_font(r, 16, True, FONT_HEADING)
    doc.add_paragraph()


def add_cover_info_table(doc: Document, rows: list[list[str]]):
    """封面学生信息表（与模板一致：左标签右内容，Table Grid）。"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for ri, (label, value) in enumerate(rows):
        c0 = table.rows[ri].cells[0]
        c1 = table.rows[ri].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, SIZE_BODY, False, FONT_BODY)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, SIZE_BODY, False, FONT_BODY)
    doc.add_paragraph()


def add_attachment_cover(doc: Document, project_line: str):
    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 1.5, 12, 6)
    r = p.add_run("附件：小组项目文档（每小组一份）")
    set_run_font(r, 14, True, FONT_HEADING)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    r2 = p2.add_run(project_line)
    set_run_font(r2, 16, True, FONT_HEADING)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("详细设计说明书")
    set_run_font(r3, 16, True, FONT_HEADING)


def add_image_figure(
    doc: Document,
    caption: str,
    image_path: Path | str,
    width_cm: float = 14.0,
    hint: str = "",
):
    """插入居中图片与图题；文件不存在时回退为占位说明。"""
    path = Path(image_path)
    if path.exists() and path.stat().st_size > 500:
        p_img = doc.add_paragraph()
        _set_paragraph_format(p_img, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
        p_img.add_run().add_picture(str(path), width=Cm(width_cm))
        p_cap = doc.add_paragraph()
        _set_paragraph_format(p_cap, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
        r = p_cap.add_run(caption)
        set_run_font(r, SIZE_BODY, True, FONT_BODY)
        doc.add_paragraph()
        return True
    add_placeholder_figure(doc, caption, hint or str(path))
    return False


def add_placeholder_figure(doc: Document, caption: str, hint: str):
    add_p(doc, f"【{caption}】", bold=True, first_line_indent=0)
    add_p(doc, hint)
    add_p(doc, "（图片占位：请从 01.需求分析 对应子目录粘贴截图或导出 PNG 后插入本文档）")


def count_doc_chars(doc: Document) -> int:
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                text += cell.text
    return len(text.replace(" ", "").replace("\n", ""))


def save_doc(doc: Document, path) -> int:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return count_doc_chars(doc)
