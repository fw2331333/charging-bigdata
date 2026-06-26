#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成需求规格说明书 Word 文档（宋体）。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

OUT = Path(r"F:\项目2资料\项目资料提交\01.需求分析\06.需求规格说明书\新能源充电桩大数据分析项目-需求规格说明书.docx")
FONT = "宋体"


def set_run_font(run, size=12, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str, level: int = 0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 22, True)
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(r, sizes.get(level, 12), True)


def add_para(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 12, bold)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, 11, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run_font(r, 11)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    add_title(doc, "新能源充电桩大数据分析项目")
    add_title(doc, "软件需求规格说明书")
    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ["文档名称", "软件需求规格说明书"],
        ["项目名称", "新能源充电桩大数据分析项目"],
        ["文档编号", "CBP-SRS-001"],
        ["版本", "V1.3"],
        ["编写日期", "2026-06-22"],
        ["适用对象", "项目组、指导教师、答辩评审"],
    ])

    add_title(doc, "1. 引言", 1)
    add_title(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档是新能源充电桩大数据分析项目的顶层需求规格说明，规定系统应具备的功能、性能、数据、接口与安全需求及验收标准，作为概要设计、详细设计、编码实现、测试与项目验收的正式依据。")

    add_title(doc, "1.2 项目范围", 2)
    add_para(doc, "本系统面向新能源汽车充电桩运营场景，构建覆盖「数据采集存储→分布式批处理→关系型持久化→BI可视化→Python分析建模→Web展示与预测→用户鉴权」的完整大数据分析应用。")
    add_para(doc, "范围内：", True)
    add_bullets(doc, [
        "HDFS集中存储四个CSV数据集（由用户上传，不入代码仓库）",
        "Hadoop MapReduce批处理统计（v1~v7）",
        "MySQL结果表与BI看板",
        "Python数据分析与机器学习预测",
        "FastAPI + Vue3 Web系统",
        "用户登录与JWT鉴权（设计已完成，见FR-08）",
    ])
    add_para(doc, "范围外：", True)
    add_bullets(doc, [
        "实时流处理（Kafka/Flink）",
        "充电桩硬件直连采集",
        "移动端App",
        "多租户SaaS计费",
    ])

    add_title(doc, "1.3 定义与缩略语", 2)
    add_table(doc, ["术语", "说明"], [
        ["HDFS", "Hadoop分布式文件系统"],
        ["MapReduce", "Hadoop批处理计算框架"],
        ["SOC", "电池荷电状态（%）"],
        ["BI", "商业智能可视化看板"],
        ["JWT", "JSON Web Token，无状态鉴权令牌"],
        ["WebHDFS", "HDFS HTTP REST访问接口"],
        ["SRS", "Software Requirements Specification"],
    ])

    add_title(doc, "2. 项目概述", 1)
    add_title(doc, "2.1 产品描述", 2)
    add_para(doc, "系统对电池运行数据与充电会话数据进行批量加工与智能分析，帮助运营方：监控电池状态、洞察充电行为规律、预测充电费用、通过Web与BI看板统一展示。")

    add_title(doc, "2.2 用户特征", 2)
    add_table(doc, ["用户角色", "特征", "主要操作"], [
        ["系统管理员", "具备Hadoop/MySQL运维能力", "集群管理、用户管理、数据上传"],
        ["数据分析师", "熟悉SQL、Python", "MR任务、入库、模型训练"],
        ["运营管理员", "业务背景，轻技术", "登录系统、查看图表与预测"],
        ["答辩评审", "外部评估", "验证功能与文档完整性"],
    ])

    add_title(doc, "2.3 运行环境", 2)
    add_title(doc, "2.3.1 目标部署环境（交付）", 3)
    add_para(doc, "三节点Hadoop分布式集群（非伪分布式）：")
    add_table(doc, ["节点", "角色"], [
        ["节点1（Master）", "NameNode、ResourceManager、MySQL、FastAPI"],
        ["节点2（Slave）", "DataNode、NodeManager"],
        ["节点3（Slave）", "DataNode、NodeManager"],
    ])
    add_title(doc, "2.3.2 开发与演示环境", 3)
    add_table(doc, ["组件", "版本要求"], [
        ["Hadoop", "3.1.x"],
        ["JDK", "1.8"],
        ["Maven", "3.6+"],
        ["MySQL", "5.7+ / 8.0"],
        ["Python", "3.11（Anaconda py3.11）"],
        ["Node.js", "18+"],
    ])

    add_title(doc, "2.4 约束与假设", 2)
    add_table(doc, ["类型", "说明"], [
        ["数据约束", "原始CSV由用户上传HDFS /Car/，禁止写入Git仓库"],
        ["技术约束", "批处理使用MapReduce；Web后端FastAPI"],
        ["网络约束", "分析节点可访问NameNode WebHDFS（9870）"],
        ["假设", "三节点集群可启动；MySQL可连接；字段与指导文档一致"],
    ])

    add_title(doc, "3. 系统功能需求", 1)
    add_title(doc, "3.1 功能结构", 2)
    add_para(doc, "系统包含八大功能模块：数据准备、批处理统计、数据持久化、BI可视化、Python分析、机器学习、Web展示、用户鉴权。")

    add_title(doc, "3.2 FR-01 数据准备模块", 2)
    add_table(doc, ["项", "内容"], [
        ["需求编号", "FR-01"],
        ["需求名称", "HDFS数据准备"],
        ["优先级", "P0"],
        ["描述", "四个CSV集中存储于HDFS /Car/"],
        ["输入", "dsv13r1、dsv13r2、nvv2t、nvv2t_md"],
        ["验收标准", "hdfs dfs -ls /Car 显示4文件；可head查看"],
    ])

    add_title(doc, "3.3 FR-02 批处理统计模块（MapReduce）", 2)
    add_para(doc, "从dsv13r1.csv完成电池状态分布式统计。最低要求≥3项；项目目标7项全部完成。")
    add_table(doc, ["子编号", "指标", "聚合键", "统计方式", "MySQL表"], [
        ["FR-02-01", "每小时平均组电压/充电电流", "yyyyMMddHH", "双指标均值", "t_voltage_current"],
        ["FR-02-02", "每小时单体电压极值", "yyyyMMddHH", "max/min", "t_cell_voltage_range"],
        ["FR-02-03", "时刻温度极值", "yyyyMMddHHmmss", "max/min", "t_temperature"],
        ["FR-02-04", "每小时能量/容量均值", "yyyyMMddHH", "双指标均值", "t_energy_capacity"],
        ["FR-02-05", "充电电流均值/峰值", "yyyyMMddHH", "均值+最大", "t_charge_current_stats"],
        ["FR-02-06", "电压-电流关系", "yyyyMMddHHmmss", "双指标均值", "t_voltage_current_relation"],
        ["FR-02-07", "SOC分段温度", "SOC区间", "双指标均值", "t_soc_temperature"],
    ])

    add_title(doc, "3.4 FR-03 数据持久化模块", 2)
    add_para(doc, "MapReduce输出导入MySQL charging_bigdata库。执行schema.sql与load_mr_to_mysql.py，支持UPSERT重复导入。")

    add_title(doc, "3.5 FR-04 BI可视化模块", 2)
    add_para(doc, "连接MySQL制作Dashboard。最低：1个Dashboard、≥3张图表、≥3个SQL View。推荐：电压电流折线、温度双折线、SOC-温度柱状图。")

    add_title(doc, "3.6 FR-05 Python/Spark 数据分析模块", 2)
    add_para(doc, "HDFS 明细经 Spark SQL 离线聚合写入 MySQL ADS 表；前端 ECharts 通过 /api/v1/bi 动态渲染。5 类分析全部完成。")
    add_table(doc, ["子编号", "分析任务", "图表", "实现", "状态"], [
        ["FR-05-01", "每日充电次数", "折线图", "Spark ADS → t_charging_daily → ECharts", "已实现"],
        ["FR-05-02", "每月充电次数", "折线图", "Spark ADS → t_charging_monthly → ECharts", "已实现"],
        ["FR-05-03", "每小时平均SOC", "折线图", "Spark ADS → t_soc_hourly → ECharts", "已实现"],
        ["FR-05-04", "SOC热力图", "热力图", "Spark ADS → t_soc_heatmap → ECharts", "已实现"],
        ["FR-05-05", "平均充电速率", "折线图", "Spark ADS → t_charge_rate_hourly → ECharts", "已实现"],
    ])

    add_title(doc, "3.7 FR-06 机器学习模块", 2)
    add_para(doc, "在分布式存储环境下完成批量样本训练与在线推理。详见《机器学习分布式批处理工作流程说明书》。")
    add_table(doc, ["子编号", "预测任务", "类型", "数据源", "目标变量", "状态"], [
        ["FR-06-01", "充电费用预测", "回归", "nvv2t_md.csv", "charging_fees", "已实现"],
        ["FR-06-02", "充电时间预测", "回归", "nvv2t_md.csv", "chargeTimeHrs", "已实现"],
        ["FR-06-03", "用户平台预测", "分类", "nvv2t_md.csv", "platform", "已实现"],
        ["FR-06-04", "剩余电量预测", "回归", "dsv13r2.csv", "soc", "已实现"],
    ])
    add_para(doc, "四项预测统一实现：", True)
    add_bullets(doc, [
        "训练脚本：analytics/scripts/train_all_models.py + analytics/fee_models.py",
        "费用算法：线性回归 vs XGBoost 对比选优（删除原 RandomForest 费用方案）",
        "模型：fee/duration/platform/soc_model.pkl",
        "评估：GET /api/v1/predict/metrics（R²、MAE、MSE、Accuracy、F1）",
        "推理：POST /api/v1/predict/fee|duration|platform|soc",
        "数据：HDFS WebHDFS流式读取或本地--local-path",
    ])

    add_title(doc, "3.8 FR-07 Web展示模块", 2)
    add_table(doc, ["子编号", "页面", "路由", "验收"], [
        ["FR-07-01", "系统首页", "/", "导航与说明"],
        ["FR-07-02", "SOC折线分析", "/soc", "展示每小时平均SOC"],
        ["FR-07-03", "SOC热力图", "/soc-heatmap", "展示日期×小时热力图"],
        ["FR-07-04", "充电次数分析", "/charging", "展示日/月充电次数"],
        ["FR-07-05", "充电速率分析", "/charge-rate", "展示平均充电速率"],
        ["FR-07-06", "四项预测", "/predict", "费用/时长/平台/SOC四个Tab"],
        ["FR-07-07", "API文档", "/docs", "Swagger可访问"],
    ])

    add_title(doc, "3.9 FR-08 用户鉴权模块", 2)
    add_para(doc, "详见《登录鉴权设计说明书》。管理端需身份认证后访问业务功能。")
    add_table(doc, ["子编号", "功能", "要求"], [
        ["FR-08-01", "用户登录", "用户名+密码，返回JWT"],
        ["FR-08-02", "密码存储", "bcrypt单向哈希，禁止明文"],
        ["FR-08-03", "会话管理", "Access Token + Refresh Token"],
        ["FR-08-04", "角色权限", "admin全权限、user查看+预测"],
        ["FR-08-05", "路由保护", "未登录访问业务API返回401"],
        ["FR-08-06", "用户表", "sys_user（auth_schema.sql）"],
    ])

    add_title(doc, "4. 数据需求", 1)
    add_title(doc, "4.1 dsv13r1.csv（MapReduce输入）", 2)
    add_para(doc, "无表头，11列。字段：id、record_time、soc、pack_voltage、charge_current、max/min_cell_voltage、max/min_temperature、available_energy、available_capacity。")
    add_title(doc, "4.2 其他数据集", 2)
    add_bullets(doc, [
        "dsv13r2.csv：有表头，Python分析（SOC、充电速率）",
        "nvv2t.csv：充电行为分析（次数、时段、平台）",
        "nvv2t_md.csv：机器学习训练（费用、时间、平台预测）",
    ])
    add_title(doc, "4.3 数据安全", 2)
    add_bullets(doc, [
        "原始CSV不得提交Git",
        "数据库密码、JWT密钥通过.env配置",
        "用户密码bcrypt存储",
    ])

    add_title(doc, "5. 外部接口需求", 1)
    add_table(doc, ["接口", "说明"], [
        ["HDFS CLI", "hdfs dfs上传与查看"],
        ["WebHDFS", "http://{nn}:9870/webhdfs/v1/..."],
        ["MySQL", "pymysql，端口3306"],
        ["REST API", "FastAPI HTTP JSON"],
        ["BI JDBC", "MySQL数据源"],
    ])

    add_title(doc, "6. 非功能需求", 1)
    add_title(doc, "6.1 性能", 2)
    add_bullets(doc, [
        "百万级CSV通过MR分布式处理",
        "Python训练支持max-rows采样",
        "API查询LIMIT≤1000",
    ])
    add_title(doc, "6.2 安全", 2)
    add_bullets(doc, [
        "密码bcrypt哈希",
        "JWT密钥不入Git",
        "SQL参数化防注入",
        "CORS白名单",
        "数据库账号最小权限",
    ])
    add_title(doc, "6.3 可扩展性", 2)
    add_bullets(doc, [
        "支持三节点分布式集群部署",
        "数据量×10时可调MR并行度、Spark ML扩展",
    ])

    add_title(doc, "7. 验收标准总表", 1)
    add_table(doc, ["模块", "最低验收", "项目目标"], [
        ["HDFS", "4文件可访问", "完成"],
        ["MapReduce", "≥3任务", "7任务"],
        ["MySQL", "≥3表有数据", "7表"],
        ["BI", "1 Dashboard、3图", "5+图"],
        ["Python分析", "2张图", "5类分析"],
        ["ML预测", "1接口可用", "4项预测+评估指标"],
        ["Web", "首页+2分析+1预测", "全页面"],
        ["鉴权", "—", "登录+JWT"],
    ])

    add_title(doc, "8. 机器学习分布式批处理工作流程（摘要）", 1)
    add_para(doc, "HDFS数据→WebHDFS流式读取→特征构造→train_all_models.py训练四个模型→joblib保存→FastAPI PredictService在线推理。")
    add_para(doc, "与MapReduce分工：MR处理dsv13r1聚合入MySQL；ML处理nvv2t_md会话级预测建模。")

    add_title(doc, "9. 登录鉴权流程（摘要）", 1)
    add_para(doc, "用户提交用户名密码 → 服务端查询sys_user → bcrypt.verify校验 → 签发JWT → 前端存储Token → 业务请求Header携带Bearer Token → 未授权返回401。")

    add_title(doc, "10. 附录", 1)
    add_title(doc, "附录A 提交清单", 2)
    add_bullets(doc, [
        "项目源码",
        "数据库建表SQL（schema.sql + auth_schema.sql）",
        "HDFS上传和运行命令说明",
        "MapReduce任务说明",
        "BI看板截图",
        "Web页面截图",
        "模型训练和评估结果",
        "项目报告",
        "小组分工说明",
        "软件需求规格说明书（本文档）",
    ])
    add_title(doc, "附录B 修订记录", 2)
    add_table(doc, ["版本", "日期", "修订内容"], [
        ["V1.0", "2026-06-22", "初稿"],
        ["V1.1", "2026-06-22", "补充鉴权、ML分布式流程、三节点集群、全功能需求"],
        ["V1.2", "2026-06-22", "依据指导文档全量落地：5类Python分析、4项ML预测、Web全页面"],
        ["V1.3", "2026-06-22", "费用预测改为线性回归vs XGBoost对比选优，删除RandomForest费用方案"],
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()
