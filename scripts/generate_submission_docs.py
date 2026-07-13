#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成项目资料提交 Word 文档（宋体），并同步 Markdown 到提交目录。"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r"F:\项目2资料")
DOCS = ROOT / "docs"
SUBMIT = ROOT / "项目资料提交"
FONT = "宋体"


def set_run_font(run, size=12, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str, level: int = 0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 22, True)
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(r, sizes.get(level, 12), True)


def add_para(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 12, bold)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, 11, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run_font(r, 11)
    doc.add_paragraph()


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    return doc


def sync_markdown():
    pairs = [
        (DOCS / "需求规格说明书.md", SUBMIT / "01.需求分析/06.需求规格说明书/新能源充电桩大数据分析项目-需求规格说明书.md"),
        (DOCS / "数据库设计说明书.md", SUBMIT / "02.系统设计/01.数据库设计/01.数据库设计文档/新能源充电桩大数据分析项目-数据库设计说明书.md"),
        (DOCS / "概要设计说明书.md", SUBMIT / "04.项目概要设计说明书/新能源充电桩大数据分析项目-概要设计说明书.md"),
    ]
    for src, dst in pairs:
        dst.parent.mkdir(parents=True, exist_ok=True)
        if src.exists():
            shutil.copy2(src, dst)
            print(f"已同步: {dst.name}")


def build_srs_docx():
    doc = new_doc()
    add_title(doc, "新能源充电桩大数据分析项目")
    add_title(doc, "软件需求规格说明书")
    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ["文档名称", "软件需求规格说明书"],
        ["项目名称", "新能源充电桩大数据分析项目"],
        ["文档编号", "CBP-SRS-001"],
        ["版本", "V1.3"],
        ["编写日期", "2026-06-22"],
    ])

    add_title(doc, "1. 引言", 1)
    add_para(doc, "本文档规定系统应具备的功能、性能、数据、接口与安全需求及完成标准，作为设计、实现与交付确认的正式依据。")
    add_para(doc, "【图片占位：功能结构图见 01.需求分析/01.功能图/】", True)

    add_title(doc, "2. 项目概述", 1)
    add_para(doc, "面向新能源汽车充电桩运营，构建 HDFS → MapReduce → MySQL → BI/Web → ML 预测完整链路。")
    add_table(doc, ["用户角色", "主要操作"], [
        ["系统管理员", "集群管理、用户管理、数据上传"],
        ["数据分析师", "MR 任务、入库、模型训练"],
        ["运营管理员", "登录、查看图表与预测"],
    ])

    add_title(doc, "3. 功能需求摘要", 1)
    add_table(doc, ["编号", "模块", "核查要点"], [
        ["FR-01", "HDFS 数据准备", "4 个 CSV 上传 /Car/"],
        ["FR-02", "MapReduce v1~v7", "7 项电池指标批处理"],
        ["FR-03", "MySQL 落库", "7 张结果表 UPSERT"],
        ["FR-04", "BI 看板", "Dashboard + SQL View"],
        ["FR-05", "Python 分析", "5 类分析图（日/月次数、SOC、热力图、速率）"],
        ["FR-06", "ML 预测", "费用/时长/平台/SOC 四项"],
        ["FR-07", "Web 展示", "Vue3 + FastAPI 全页面"],
        ["FR-08", "用户鉴权", "邮箱登录 + JWT + bcrypt"],
    ])

    add_title(doc, "4. 部署环境（实际）", 1)
    add_bullets(doc, [
        "Hadoop：三节点分布式集群（课程交付目标）",
        "云演示：阿里云 ECS Ubuntu 22.04，Docker Compose 部署",
        "访问地址：http://115.29.194.137:8080",
        "演示账号：admin@example.com / admin123",
    ])

    add_title(doc, "5. 完成标准", 1)
    add_table(doc, ["模块", "最低", "本项目"], [
        ["MapReduce", "≥3 任务", "7 任务"],
        ["MySQL", "≥3 表", "7 表 + ADS 5 表"],
        ["Python 分析", "2 张图", "5 类"],
        ["ML", "1 接口", "4 项预测"],
        ["Web", "首页+分析+预测", "全页面+鉴权"],
    ])

    add_title(doc, "6. 修订记录", 1)
    add_table(doc, ["版本", "日期", "说明"], [
        ["V1.3", "2026-06-22", "费用预测 XGBoost 选优；全功能落地"],
    ])

    out = SUBMIT / "01.需求分析/06.需求规格说明书/新能源充电桩大数据分析项目-需求规格说明书.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"已生成: {out}")


def build_db_docx():
    doc = new_doc()
    add_title(doc, "新能源充电桩大数据分析项目")
    add_title(doc, "数据库设计说明书")
    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ["数据库名", "charging_bigdata"],
        ["字符集", "utf8mb4"],
        ["版本", "V1.2"],
    ])

    add_title(doc, "1. 设计概述", 1)
    add_para(doc, "MySQL 存储 MapReduce 聚合结果、ADS 分析汇总及用户认证数据；原始 CSV 仅存 HDFS。")
    add_para(doc, "【图片占位：E-R 图见 02.系统设计/01.数据库设计/02.E-R图/】", True)

    add_title(doc, "2. MR 结果表清单", 1)
    add_table(doc, ["表名", "MR", "维度", "用途"], [
        ["t_voltage_current", "v1", "小时", "电压电流折线"],
        ["t_cell_voltage_range", "v2", "小时", "单体电压极值"],
        ["t_temperature", "v3", "时间点", "温度双折线"],
        ["t_energy_capacity", "v4", "小时", "能量容量"],
        ["t_charge_current_stats", "v5", "小时", "充电电流"],
        ["t_voltage_current_relation", "v6", "时间点", "电压电流关系"],
        ["t_soc_temperature", "v7", "SOC 段", "SOC-温度柱状"],
    ])

    add_title(doc, "3. ADS 汇总表", 1)
    add_table(doc, ["表名", "说明"], [
        ["t_soc_hourly", "每小时平均 SOC"],
        ["t_charging_daily", "每日充电次数"],
        ["t_charging_monthly", "每月充电次数"],
        ["t_charge_rate_hourly", "每小时充电速率"],
        ["t_soc_heatmap", "SOC 热力图"],
    ])

    add_title(doc, "4. 认证表", 1)
    add_bullets(doc, [
        "sys_user：用户邮箱、bcrypt 密码、角色",
        "sys_refresh_token：双 Token 刷新",
        "sys_email_verification_token：注册/重置密码邮件",
    ])

    add_title(doc, "5. BI 视图", 1)
    add_bullets(doc, ["v_voltage_current", "v_temperature", "v_soc_temperature"])

    add_title(doc, "6. 入库策略", 1)
    add_para(doc, "load_mr_to_mysql.py 读取 HDFS part-r-00000，按唯一键 ON DUPLICATE KEY UPDATE。")

    out = SUBMIT / "02.系统设计/01.数据库设计/01.数据库设计文档/新能源充电桩大数据分析项目-数据库设计说明书.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"已生成: {out}")


def build_outline_docx():
    doc = new_doc()
    add_title(doc, "新能源充电桩大数据分析项目")
    add_title(doc, "概要设计说明书")
    doc.add_paragraph()

    add_title(doc, "1. 总体架构", 1)
    add_para(doc, "五层架构：展示层（Vue3/BI）→ 应用层（FastAPI）→ 分析层（Python/ML）→ 计算层（MapReduce）→ 存储层（HDFS/MySQL）。")
    add_para(doc, "【图片占位：系统架构图见 01.需求分析/01.功能图/】", True)

    add_title(doc, "2. 技术选型", 1)
    add_table(doc, ["层次", "技术"], [
        ["分布式存储", "Hadoop HDFS 3.x"],
        ["批处理", "MapReduce Java 8"],
        ["数据库", "MySQL 8.0"],
        ["后端", "FastAPI + Uvicorn"],
        ["前端", "Vue3 + TypeScript + ECharts"],
        ["机器学习", "scikit-learn + XGBoost"],
        ["部署", "Docker Compose（MySQL + Backend + Nginx）"],
    ])

    add_title(doc, "3. 模块划分", 1)
    add_table(doc, ["模块", "目录", "职责"], [
        ["MapReduce", "mapreduce/", "v1~v7 批处理"],
        ["数据分析", "analytics/", "HDFS 读取、训练、绘图"],
        ["后端", "backend/", "REST API、鉴权、预测"],
        ["前端", "frontend/", "页面与图表"],
        ["SQL", "sql/", "建表与种子"],
    ])

    add_title(doc, "4. 部署架构（实际）", 1)
    add_bullets(doc, [
        "课程目标：三节点 Hadoop 集群",
        "现场演示：阿里云 ECS + Docker，前端 8080 端口",
        "Nginx 反代 /api → backend:8000",
    ])

    add_title(doc, "5. 安全设计", 1)
    add_bullets(doc, [
        "JWT 双 Token + bcrypt 密码",
        "业务 API 需 Bearer 认证",
        "敏感配置 .env.docker，不入 Git",
    ])

    out = SUBMIT / "04.项目概要设计说明书/新能源充电桩大数据分析项目-概要设计说明书.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"已生成: {out}")


def main():
    sync_markdown()
    build_srs_docx()
    build_db_docx()
    build_outline_docx()
    print("全部完成。")


if __name__ == "__main__":
    main()
